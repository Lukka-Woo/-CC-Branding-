import sys, os, io

_BRAND   = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
_PROJECT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, _BRAND)

_DOCS  = os.path.join(_PROJECT, 'docs')
_MEDIA = os.path.join(_PROJECT, 'media')

os.makedirs(_DOCS, exist_ok=True)

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from scripts.docx_builder import (
    BrandDocx, _set_font, _cell_shading, _cell_borders,
    _para_left_bar, _para_shading,
    _GRAY, _GREEN, _NONE,
)
import scripts.brand_tokens as BT

# ── Font constants ─────────────────────────────────────────────────────────────
_CN_FONT = 'Alibaba PuHuiTi 2.0'
_EN_FONT = 'Inter'

# ── Font helpers ───────────────────────────────────────────────────────────────

def _set_run_font(run, font_cn=_CN_FONT, font_en=_EN_FONT):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),     font_en)
    rFonts.set(qn('w:hAnsi'),     font_en)
    rFonts.set(qn('w:eastAsia'),  font_cn)
    rFonts.set(qn('w:cs'),        font_cn)
    run.font.name = font_en


def _set_doc_font(raw_doc):
    """Apply Alibaba PuHuiTi 2.0 as the document default East-Asian font."""
    styles_elem = raw_doc.styles._element
    # 1. docDefaults in styles.xml
    docDefs = styles_elem.find(qn('w:docDefaults'))
    if docDefs is not None:
        rPrDef = docDefs.find(qn('w:rPrDefault'))
        if rPrDef is None:
            rPrDef = OxmlElement('w:rPrDefault')
            docDefs.append(rPrDef)
        rPr = rPrDef.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDef.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'),    _EN_FONT)
        rFonts.set(qn('w:hAnsi'),    _EN_FONT)
        rFonts.set(qn('w:eastAsia'), _CN_FONT)
        rFonts.set(qn('w:cs'),       _CN_FONT)

    # 2. Normal style
    try:
        normal = raw_doc.styles['Normal']
        normal.font.name = _EN_FONT
        rPr2 = normal.font._element
        rFonts2 = rPr2.find(qn('w:rFonts'))
        if rFonts2 is None:
            rFonts2 = OxmlElement('w:rFonts')
            rPr2.insert(0, rFonts2)
        rFonts2.set(qn('w:eastAsia'), _CN_FONT)
    except Exception:
        pass

    # 3. All runs in body paragraphs
    for para in raw_doc.paragraphs:
        for run in para.runs:
            _set_run_font(run)

    # 4. All runs in body tables
    for tbl in raw_doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run)


def _inject_saved_header(docx_path, src_docx_path):
    """Copy header1.xml and its rels from src_docx into docx_path using zip surgery."""
    import zipfile, shutil, tempfile

    with zipfile.ZipFile(src_docx_path, 'r') as src_z:
        hdr_bytes  = src_z.read('word/header1.xml')
        hdr_rels   = src_z.read('word/_rels/header1.xml.rels')

    # Rewrite the destination DOCX via temp file
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp,      'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == 'word/header1.xml':
                zout.writestr(item, hdr_bytes)
            elif item.filename == 'word/_rels/header1.xml.rels':
                zout.writestr(item, hdr_rels)
            else:
                zout.writestr(item, zin.read(item.filename))

    shutil.move(tmp, docx_path)


# ── Content helpers ────────────────────────────────────────────────────────────

def _bullet(raw_doc, text, level=0, sz=10.5):
    p = raw_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    indent = Pt(16 + level * 14)
    p.paragraph_format.left_indent       = indent
    p.paragraph_format.first_line_indent = Pt(-13)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.35
    bullet_char = '•' if level == 0 else '–'
    r1 = p.add_run(f'{bullet_char}  ')
    _set_font(r1, size_pt=sz, color_hex=BT.PRIMARY_500_HEX, bold=True)
    _set_run_font(r1)
    r2 = p.add_run(text)
    _set_font(r2, size_pt=sz, color_hex=BT.NEUTRAL_700_HEX)
    _set_run_font(r2)
    return p


def _cover(raw_doc, vesuvius_logo_path, our_logo_png_path):
    # Green top accent strip
    accent = raw_doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(0)
    pPr = accent._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), BT.PRIMARY_500_HEX.lstrip('#'))
    pPr.append(shd)
    r = accent.add_run('  ')
    r.font.size = Pt(6)

    for _ in range(3):
        sp = raw_doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Vesuvius logo — large, centered
    if os.path.exists(vesuvius_logo_path):
        lp = raw_doc.add_paragraph()
        lp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(4)
        lp.add_run().add_picture(vesuvius_logo_path, width=Mm(52))

    # "致 Vesuvius 中国区" tag
    tag_p = raw_doc.add_paragraph()
    tag_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_after  = Pt(0)
    tr = tag_p.add_run('致 Vesuvius 中国区  |  PROPOSAL FOR VESUVIUS CHINA')
    _set_font(tr, size_pt=8.5, color_hex=BT.NEUTRAL_400_HEX)
    _set_run_font(tr)

    # Divider rule
    rule_p = raw_doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(14)
    rule_p.paragraph_format.space_after  = Pt(14)
    pPr2 = rule_p._p.get_or_add_pPr()
    bdr2 = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BT.PRIMARY_500_HEX.lstrip('#'))
    bdr2.append(bot)
    pPr2.append(bdr2)

    # Main title
    title_p = raw_doc.add_paragraph()
    title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after  = Pt(6)
    tr2 = title_p.add_run('企业培训平台功能增强建设建议书')
    _set_font(tr2, size_pt=24, bold=True, color_hex=BT.NEUTRAL_900_HEX)
    _set_run_font(tr2)

    en_p = raw_doc.add_paragraph()
    en_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    en_p.paragraph_format.space_before = Pt(0)
    en_p.paragraph_format.space_after  = Pt(16)
    er = en_p.add_run('Enterprise Training Platform Enhancement Proposal')
    _set_font(er, size_pt=11, color_hex=BT.NEUTRAL_400_HEX)
    _set_run_font(er)

    # Scope summary
    scope_p = raw_doc.add_paragraph()
    scope_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    scope_p.paragraph_format.space_after  = Pt(4)
    scope_r = scope_p.add_run(
        '培训内容管理  ·  课程部署  ·  积分激励  ·  视频培训  ·  培训矩阵  ·  数据台账'
    )
    _set_font(scope_r, size_pt=10, color_hex=BT.NEUTRAL_700_HEX)
    _set_run_font(scope_r)

    for _ in range(4):
        raw_doc.add_paragraph()

    # Bottom meta row
    meta_tbl = raw_doc.add_table(1, 3)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in meta_tbl.rows[0].cells:
        _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

    lc = meta_tbl.rows[0].cells[0]
    lc.width = Mm(55)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp2 = lc.paragraphs[0]
    lp2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp2.paragraph_format.space_before = Pt(0)
    lp2.paragraph_format.space_after  = Pt(0)
    if os.path.exists(our_logo_png_path):
        lp2.add_run().add_picture(our_logo_png_path, width=Mm(35))

    cc = meta_tbl.rows[0].cells[1]
    cc.width = Mm(55)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cc.paragraphs[0]
    cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(2)
    cr = cp.add_run(BT.BRAND_FULL_CN)
    _set_font(cr, size_pt=8, color_hex=BT.NEUTRAL_700_HEX)
    _set_run_font(cr)

    rc = meta_tbl.rows[0].cells[2]
    rc.width = Mm(45)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(2)
    rr = rp.add_run('V2.0  ·  2026 年 6 月\n机密文件，仅供内部参阅')
    _set_font(rr, size_pt=7.5, color_hex=BT.NEUTRAL_400_HEX, italic=True)
    _set_run_font(rr)

    # Page break
    pb = raw_doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(0)
    pb.add_run().add_break(
        __import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE
    )


# ══════════════════════════════════════════════════════════════════════════════
# Build document
# ══════════════════════════════════════════════════════════════════════════════

doc     = BrandDocx(doc_type='解决方案建议书')
raw_doc = doc.document

sec = raw_doc.sections[0]
sec.different_first_page_header_footer = True

# Fix footer disclaimer
for tbl in sec.footer.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text and ('人力资源' in run.text or '证明使用' in run.text):
                        run.text = '机密文件 · 仅供 Vesuvius 中国区内部参阅'

# Cover page
_VESUVIUS_LOGO = os.path.join(_MEDIA, 'covers', 'VESUVIUS_logo.png')
_OUR_LOGO      = BT.LOGO_HORIZONTAL_PRIMARY_PNG
_cover(raw_doc, _VESUVIUS_LOGO, _OUR_LOGO)

# ─────────────────────────────────────────────────────────────────────────────
# 正文
# ─────────────────────────────────────────────────────────────────────────────

doc.add_title(
    '企业培训平台功能增强建设建议书',
    subtitle='ENTERPRISE TRAINING PLATFORM ENHANCEMENT PROPOSAL  ·  FOR VESUVIUS CHINA',
)

doc.add_info_table([
    ('客户单位',   'Vesuvius 中国区'),
    ('项目名称',   '企业培训平台功能增强建设项目'),
    ('建议版本',   'V2.0'),
    ('预计周期',   '约 8 周（视最终范围确认）'),
], label_width_mm=35, value_width_mm=110)

# ── 一、项目背景与目标 ─────────────────────────────────────────────────────────
doc.add_heading('一、项目背景与目标', level=1)
doc.add_body(
    '贵司现有培训平台已具备基础的内容管理能力。本次增强的核心目标是让培训从"内容已发布"'
    '真正走向"任务被完成、过程有记录、结果可追溯"——让管理者不再依赖手工 Excel 和'
    '微信群催促，让每一次培训都有闭环。'
)
doc.add_body('本期建设将在六个维度产生直接的管理改变：')

_bullet(raw_doc, '培训任务有人认领、有截止时间、完成情况实时可查，执行率可量化。')
_bullet(raw_doc, '员工参与有正向激励和逾期约束，学习行为有数据记录。')
_bullet(raw_doc, '操作规程、安全演示等视频内容可上传课程并分发部署，观看完成情况自动留痕。')
_bullet(raw_doc, '一张培训矩阵实时展示所有员工的课程覆盖情况，无需人工汇总。')
_bullet(raw_doc, '学习过程真实性有基础保障，培训记录具备可信度与审计能力。')
_bullet(raw_doc, '培训记录、积分流水、考试台账具备向合规档案扩展的数据基础。')

doc.add_note(
    '建设主线：内容可管理 → 任务可部署 → 过程可监督 → 结果可统计 → 数据可导出。',
    label='一句话概括  ',
)

# ── 二、本期建设范围 ──────────────────────────────────────────────────────────
doc.add_heading('二、本期建设范围', level=1)
doc.add_body(
    '本期围绕七个功能版块展开建设，每个版块都直接对应贵司管理团队的一个具体痛点。'
)

doc.add_data_table(
    headers=['功能版块', '对 Vesuvius 的直接价值'],
    data_rows=[
        ['培训内容管理', '课程内容有人审核、有版本记录，确保员工看到的是已授权的最新版本'],
        ['课程部署',     '每次培训都有明确责任人、截止时间和完成追踪，彻底告别微信群催促'],
        ['正负积分体系', '完成课程有奖励、逾期有约束，积分记录透明，驱动员工主动学习'],
        ['视频培训',     '上传操作规程或安全演示视频，员工在线观看，系统自动记录完成状态'],
        ['培训矩阵',     '一张实时视图展示所有员工的课程完成率，部门/岗位缺口一眼可见'],
        ['学习过程监督', '系统记录学习行为真实性，减少挂机和无效观看，保障培训数据可信度'],
        ['数据统计与台账', '完成报表、考试记录、积分明细一键导出，支持内部汇报与合规备查'],
    ],
    col_widths_mm=[38, 117],
)

# ── 三、功能方案 ──────────────────────────────────────────────────────────────
doc.add_heading('三、功能方案', level=1)

# 3.1
doc.add_heading('3.1  培训内容管理', level=2)
doc.add_body(
    '管理员可对课程内容持续维护，设置不同类型课程的完成规则，并通过审核流确保每门课程'
    '在正式发布前经过负责人确认。制度宣导类课程可设为"阅读确认即完成"；安全培训类课程'
    '可叠加"视频观看比例 + 考试通过"；版本记录确保每次改动都有据可查。'
)
doc.add_data_table(
    headers=['功能', '您将获得'],
    data_rows=[
        ['灵活的完成规则',   '按课程类型分别配置：阅读确认 / 视频观看比例 / 考试通过，不强迫所有课程走同一套标准'],
        ['内容审核流',       '课程正式发布前须经培训负责人审核，防止未授权内容失误发布到平台内部'],
        ['版本记录',         '每次发布形成版本留痕，历史版本可追溯，迎检时有完整内容变更依据'],
        ['AI 辅助',          '支持 AI 辅助生成知识点新建章节，考试题目支持 AI 辅助生成，大幅降低题库维护工作量'],
    ],
    col_widths_mm=[40, 115],
)

# 3.2
doc.add_heading('3.2  课程部署', level=2)
doc.add_body(
    '课程部署是让培训从"资源库"变成"执行任务"的关键。管理员按部门、岗位或人员维度'
    '分配课程，设置截止时间和提醒规则后，系统自动追踪每个人的完成进度——谁完成了、'
    '谁逾期了、考试结果怎样，全部实时可查，无需手工汇总。'
)
doc.add_data_table(
    headers=['功能', '您将获得'],
    data_rows=[
        ['灵活的分配方式',   '支持按员工、部门、岗位、职能、班组多维度分配，一次操作覆盖目标人群'],
        ['截止时间与提醒',   '系统在临近截止时自动提醒未完成员工，逾期情况同步通知管理者'],
        ['实时完成追踪',     '管理员随时查看当前部署的完成率、未完成人员清单与考试通过情况'],
        ['部署历史存档',     '每次培训的对象、版本、完成情况长期留存，支持年度培训复盘与审查'],
    ],
    col_widths_mm=[40, 115],
)

# 3.3
doc.add_heading('3.3  正负积分体系', level=2)
doc.add_body(
    '积分体系以数据化方式建立轻量级激励和约束，代替反复的行政催促。完成课程自动加分，'
    '逾期自动扣分，每一笔积分流水都有明细可查——员工看得到自己的表现，管理者看得到'
    '团队的整体参与度。后续如有需要，积分数据可直接扩展为学习排行榜或年度激励评选的依据。'
)
doc.add_data_table(
    headers=['积分规则', '员工视角', '管理者视角'],
    data_rows=[
        ['完成课程 / 考试通过', '个人积分实时增加，可查看积分明细', '整体参与度提升，无需反复催促'],
        ['课程逾期 / 考试未过', '收到系统提示，清楚了解扣分原因', '逾期记录可导出，支持部门绩效参考'],
        ['活动奖励 / 人工调整', '积分变动有通知，透明公正',       '管理端可补发或纠错，全程留审计日志'],
    ],
    col_widths_mm=[45, 60, 50],
)

# 3.4
doc.add_heading('3.4  视频培训', level=2)
doc.add_body(
    '管理员可在课程章节中上传操作规程、安全培训演示或设备维护视频，员工通过 PC 或'
    '手机浏览器在线观看。系统实时记录每位员工的观看进度，可配置"观看达到指定比例'
    '方可完成课程"，确保视频培训留有真实记录而非仅标注已观看。'
)
doc.add_data_table(
    headers=['功能', '您将获得'],
    data_rows=[
        ['在线上传与播放',   '视频直接上传至课程，员工无需下载，PC/手机浏览器均可播放'],
        ['断点续播',         '员工上次观看到哪里，下次自动从断点继续，不影响碎片化学习'],
        ['观看进度记录',     '系统记录实际观看比例，课程完成判断基于真实数据而非点击记录'],
        ['播放控制',         '可配置是否允许拖拽进度条，防止员工跳过关键内容直接标记完成'],
    ],
    col_widths_mm=[40, 115],
)

# 3.5
doc.add_heading('3.5  培训矩阵', level=2)
doc.add_body(
    '培训矩阵解决的是"每次想了解培训覆盖情况，都要从头拉数据、手工处理 Excel"的问题。'
    '它是一张始终在线、实时更新的全员培训视图，管理者可以按部门、岗位或培训状态快速'
    '筛选，一眼定位需要关注的人员或课程，并支持一键导出。'
)
doc.add_data_table(
    headers=['视角', '可以看到什么', '典型场景'],
    data_rows=[
        ['员工个人',   '已完成哪些课程、还有哪些未完成、当前积分',    '员工自主了解学习任务进度'],
        ['部门负责人', '本部门完成率、逾期人数、课程通过率',           '月度培训汇报，无需手工统计'],
        ['培训管理者', '跨部门/岗位的完成缺口、课程执行效果对比',      '安全检查前快速准备培训台账'],
    ],
    col_widths_mm=[28, 82, 45],
)

# 3.6
doc.add_heading('3.6  学习过程监督', level=2)
doc.add_body(
    '系统通过轻量化机制提升学习的真实性：长时间无操作时弹出确认提示，页面切出时记录'
    '次数，视频进度不允许跳过未观看部分。所有过程行为均有记录，管理员可查看每位员工'
    '的学习时长和操作情况。这些机制不涉及摄像头或隐私采集，仅从行为层面保障培训数据'
    '的可信度与审计价值。'
)
doc.add_note(
    '培训记录的可信度越高，其作为安全合规台账或 ESG 培训数据的证明力越强。',
    label='管理价值  ',
)

# 3.7
doc.add_heading('3.7  数据统计与台账', level=2)
doc.add_body(
    '平台将课程完成、考试结果、积分变动、视频观看记录统一汇聚，支持按员工、部门、'
    '岗位或时间段多维度导出。无论是月度内部汇报、年度安全培训统计，还是迎接政府'
    '安全检查或 ISO 审核，都能在几分钟内生成一份完整的数字台账。'
)
doc.add_data_table(
    headers=['报表类型', '典型用途'],
    data_rows=[
        ['培训完成报表',   '日常执行管理；按部门/岗位汇总完成率，快速定位缺口'],
        ['考试结果报表',   '评估学习效果；识别需要补训的员工'],
        ['积分明细报表',   '积分核查与管理复核；支持绩效参考数据导出'],
        ['视频学习记录',   '视频课程完成判断依据；满足合规留痕要求'],
        ['培训台账导出',   '迎检备查；ESG 培训覆盖数据出口；年度培训统计'],
    ],
    col_widths_mm=[42, 113],
)

# ── 四、双方团队分工 ──────────────────────────────────────────────────────────
doc.add_heading('四、双方团队分工', level=1)

# 4.1
doc.add_heading('4.1  任务清单与工作界面', level=2)
doc.add_body(
    '项目执行期间，双方团队在各阶段均有明确的责任边界。以下任务清单列出了每项工作的分工，'
    '确保推进有序、不产生责任盲区。'
)
doc.add_data_table(
    headers=['编号', '任务描述', '甲方（Vesuvius）职责', '我方职责'],
    data_rows=[
        ['1', '需求确认',
         '确认功能范围与角色权限；指定培训负责人作为项目对接人；审批我方给予的完整需求清单',
         '输出需求清单、主要页面原型及双方认可的验收标准'],
        ['2', '数据与组织准备',
         '提供员工花名册及部门/岗位层级数据；确认课程分类体系与初始管理权限分配',
         '完成组织架构导入、账号初始化及角色权限体系配置'],
        ['3', '功能开发',
         '提供部分课程内容（文档/视频/其他格式）；对原型方案进行业务评审',
         '开发七大功能模块，每周固定进度同步，问题当周响应'],
        ['4', '系统内测',
         '组织 HR 及培训管理员参与内测，提供真实场景操作反馈',
         '根据反馈修复问题并完成回归测试，确保功能符合验收标准'],
        ['5', '上线验收',
         '组织管理员正式验收，确认系统满足建设目标后签署验收确认',
         '提供管理员培训、学员端操作手册及完整交付文件'],
        ['6', '上线后运维',
         '指定系统管理员负责日常运营；统一收集并反馈优化需求',
         '提供上线后 30 天免费维护，功能问题响应时间不超过 1 个工作日'],
    ],
    col_widths_mm=[11, 28, 65, 51],
)

# 4.2
doc.add_heading('4.2  初步项目进度安排', level=2)
doc.add_body('标准版初步建设周期约 8 周，按五个阶段有序推进：')
doc.add_data_table(
    headers=['阶段', '周次', '涉及职能', '里程碑 / 交付物'],
    data_rows=[
        ['需求确认',   '第 1 周',   '产品经理 · 甲方培训负责人',     '需求清单、权限矩阵、验收口径'],
        ['原型设计',   '第 2 周',   '产品经理 · UI 设计师',          '主要页面原型，双方评审确认'],
        ['功能开发',   '第 3–6 周', '前端 · 后端 · 数据工程师',      '七大功能模块可联调版本'],
        ['测试与修复', '第 7 周',   '测试工程师 · 甲方内测团队',     '测试记录、修复清单、权限与报表验证'],
        ['上线验收',   '第 8 周',   '产品经理 · 实施顾问',           '验收版本 · 管理员培训 · 操作手册'],
    ],
    col_widths_mm=[22, 16, 56, 61],
)

doc.add_spacer(1)
doc.add_note(
    '我们已为 Vesuvius 中国区预留了专项资源，欢迎随时安排平台演示或技术对话。'
    '期待与贵团队开始这段合作。',
    label='联系我们  ',
)
doc.add_spacer(2)
doc.add_signature_block(company=BT.BRAND_FULL_CN, show_seal=True)

# ─────────────────────────────────────────────────────────────────────────────
# 后处理：1. 应用字体  2. 注入用户修改的页眉
# ─────────────────────────────────────────────────────────────────────────────
_set_doc_font(raw_doc)

OUTPUT   = os.path.join(_DOCS, 'Vesuvius_培训平台增强建议书_v3.0.docx')
V1_DOCX  = os.path.join(_DOCS, 'Vesuvius_培训平台增强建议书_v1.0.docx')

doc.save(OUTPUT)

# Restore user's manually-edited header (zip-level surgery)
_inject_saved_header(OUTPUT, V1_DOCX)

print(f'✓ Saved: {OUTPUT}')
