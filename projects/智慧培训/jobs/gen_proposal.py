import sys, os, io

_BRAND   = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
_PROJECT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, _BRAND)

_DOCS  = os.path.join(_PROJECT, 'docs')
_MEDIA = os.path.join(_PROJECT, 'media')
_REFS  = os.path.join(_PROJECT, 'references')

os.makedirs(_DOCS, exist_ok=True)
os.makedirs(os.path.join(_DOCS, 'img'), exist_ok=True)

from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

from scripts.docx_builder import (
    BrandDocx, _set_font, _cell_shading, _cell_borders,
    _para_left_bar, _para_shading, _rgb,
    _GRAY, _GREEN, _NONE,
)
import scripts.brand_tokens as BT


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_page_break(raw_doc):
    p = raw_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)


def _add_cover(raw_doc, vesuvius_logo_path, our_logo_png_path):
    # ── Top accent bar ────────────────────────────────────────────────────────
    accent = raw_doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(0)
    from docx.oxml import OxmlElement as OE
    from docx.oxml.ns import qn as QN
    pPr = accent._p.get_or_add_pPr()
    shd = OE('w:shd')
    shd.set(QN('w:val'), 'clear')
    shd.set(QN('w:color'), 'auto')
    shd.set(QN('w:fill'), BT.PRIMARY_500_HEX.lstrip('#'))
    pPr.append(shd)
    bdr = OE('w:pBdr')
    pPr.append(bdr)
    r = accent.add_run('  ')
    r.font.size = Pt(6)

    # ── Spacer ────────────────────────────────────────────────────────────────
    for _ in range(3):
        sp = raw_doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # ── Vesuvius logo (centered, large) ───────────────────────────────────────
    if os.path.exists(vesuvius_logo_path):
        logo_p = raw_doc.add_paragraph()
        logo_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after  = Pt(4)
        lr = logo_p.add_run()
        lr.add_picture(vesuvius_logo_path, width=Mm(52))

    # ── "致 Vesuvius 中国区" tag ───────────────────────────────────────────────
    tag_p = raw_doc.add_paragraph()
    tag_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_before = Pt(4)
    tag_p.paragraph_format.space_after  = Pt(0)
    tag_r = tag_p.add_run('致 Vesuvius 中国区  |  PROPOSAL FOR VESUVIUS CHINA')
    _set_font(tag_r, size_pt=8.5, color_hex=BT.NEUTRAL_400_HEX)

    # ── Divider rule ──────────────────────────────────────────────────────────
    rule_p = raw_doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(14)
    rule_p.paragraph_format.space_after  = Pt(14)
    pPr2 = rule_p._p.get_or_add_pPr()
    bdr2 = OE('w:pBdr')
    bot  = OE('w:bottom')
    bot.set(QN('w:val'),   'single')
    bot.set(QN('w:sz'),    '12')
    bot.set(QN('w:space'), '1')
    bot.set(QN('w:color'), BT.PRIMARY_500_HEX.lstrip('#'))
    bdr2.append(bot)
    pPr2.append(bdr2)

    # ── Main title ────────────────────────────────────────────────────────────
    title_p = raw_doc.add_paragraph()
    title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after  = Pt(6)
    tr = title_p.add_run('企业培训平台功能增强建设建议书')
    _set_font(tr, size_pt=24, bold=True, color_hex=BT.NEUTRAL_900_HEX)

    en_p = raw_doc.add_paragraph()
    en_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    en_p.paragraph_format.space_before = Pt(0)
    en_p.paragraph_format.space_after  = Pt(16)
    er = en_p.add_run('Enterprise Training Platform Enhancement Proposal')
    _set_font(er, size_pt=11, color_hex=BT.NEUTRAL_400_HEX)

    # ── Scope summary block ───────────────────────────────────────────────────
    scope_p = raw_doc.add_paragraph()
    scope_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    scope_p.paragraph_format.space_before = Pt(0)
    scope_p.paragraph_format.space_after  = Pt(4)
    scope_r = scope_p.add_run(
        '本期建设：培训内容管理  ·  课程部署  ·  正负积分激励\n'
        '视频化培训  ·  培训矩阵可视化  ·  学习过程防挂机  ·  数据台账'
    )
    _set_font(scope_r, size_pt=10, color_hex=BT.NEUTRAL_700_HEX)

    # ── Bottom meta block ─────────────────────────────────────────────────────
    for _ in range(4):
        raw_doc.add_paragraph()

    meta_tbl = raw_doc.add_table(1, 3)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in meta_tbl.rows[0].cells:
        _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

    # Left: our logo
    lc = meta_tbl.rows[0].cells[0]
    lc.width = Mm(55)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    if os.path.exists(our_logo_png_path):
        lr2 = lp.add_run()
        lr2.add_picture(our_logo_png_path, width=Mm(35))

    # Center: company name
    cc = meta_tbl.rows[0].cells[1]
    cc.width = Mm(55)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cc.paragraphs[0]
    cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(2)
    cr = cp.add_run(BT.BRAND_FULL_CN)
    _set_font(cr, size_pt=8, color_hex=BT.NEUTRAL_700_HEX)

    # Right: version + date
    rc = meta_tbl.rows[0].cells[2]
    rc.width = Mm(45)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(2)
    rr = rp.add_run('V1.0  ·  2026 年 6 月\n机密文件，仅供内部参阅')
    _set_font(rr, size_pt=7.5, color_hex=BT.NEUTRAL_400_HEX, italic=True)

    # ── Page break ────────────────────────────────────────────────────────────
    _add_page_break(raw_doc)


def _bullet(raw_doc, text, level=0, sz=11):
    p = raw_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    indent = Pt(18 + level * 14)
    p.paragraph_format.left_indent       = indent
    p.paragraph_format.first_line_indent = Pt(-14)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.4
    bullet_char = '•' if level == 0 else '–'
    r1 = p.add_run(f'{bullet_char}  ')
    _set_font(r1, size_pt=sz, color_hex=BT.PRIMARY_500_HEX, bold=True)
    r2 = p.add_run(text)
    _set_font(r2, size_pt=sz, color_hex=BT.NEUTRAL_700_HEX)
    return p


def _add_key_callout(doc, text, label='核心价值'):
    doc.add_note(text, label=f'{label}  ')


# ── Build document ────────────────────────────────────────────────────────────

doc     = BrandDocx(doc_type='解决方案建议书')
raw_doc = doc.document

# Enable different first-page header so cover has no brand header
sec = raw_doc.sections[0]
sec.different_first_page_header_footer = True

# Fix footer: replace HR disclaimer with proposal disclaimer
for tbl in sec.footer.tables:
    for row in tbl.rows:
        rc = row.cells[1]
        for para in rc.paragraphs:
            for run in para.runs:
                if run.text and ('人力资源' in run.text or '证明使用' in run.text):
                    run.text = '机密文件 · 仅供 Vesuvius 中国区内部参阅'

# Cover page
_VESUVIUS_LOGO = os.path.join(_MEDIA, 'covers', 'VESUVIUS_logo.png')
_OUR_LOGO      = BT.LOGO_HORIZONTAL_PRIMARY_PNG
_add_cover(raw_doc, _VESUVIUS_LOGO, _OUR_LOGO)

# ══════════════════════════════════════════════════════════════════════════════
# 正文
# ══════════════════════════════════════════════════════════════════════════════

doc.add_title(
    '企业培训平台功能增强建设建议书',
    subtitle='ENTERPRISE TRAINING PLATFORM ENHANCEMENT PROPOSAL  ·  FOR VESUVIUS CHINA',
)

# ── 前言 ──────────────────────────────────────────────────────────────────────
doc.add_heading('前言', level=1)
doc.add_body(
    'Vesuvius 作为全球领先的耐火材料与铸造流动控制技术企业，其中国区业务横跨多个工厂与职能体系，对'
    '员工岗位技能、安全合规及操作规程的培训管理提出了较高要求。我们理解，贵司现有培训平台已具备一定'
    '的基础能力，而本次建设的核心目标在于：让培训从"内容已发布"真正走向"任务被完成、过程有记录、'
    '结果可追溯"。'
)
doc.add_body(
    '本建议书是我们在深入理解贵司业务场景后，围绕七个功能版块制定的增强方案。我们希望通过清晰的功能'
    '边界、务实的实施路径和可量化的交付标准，与 Vesuvius 中国区建立一段长期可信赖的合作关系。'
)

doc.add_info_table([
    ('客户单位',   'Vesuvius 中国区'),
    ('项目名称',   '企业培训平台功能增强建设项目'),
    ('建议版本',   'V1.0'),
    ('建设周期',   '约 8 周（视最终范围确认）'),
    ('文件状态',   '机密 · 仅供内部参阅'),
], label_width_mm=35, value_width_mm=110)

# ── 一、项目理解与建设目标 ────────────────────────────────────────────────────
doc.add_heading('一、项目理解与建设目标', level=1)
doc.add_body(
    '本项目是在贵司现有培训平台基础上的功能增强，而非推倒重建。我们的目标是补齐管理闭环中关键的'
    '"最后一公里"——让已沉淀的培训内容真正转化为可执行的学习任务，让分散在各处的学习记录汇聚成'
    '管理者可以依赖的数字台账。'
)
doc.add_body(
    '具体而言，本期建设将围绕以下六个方向产生可量化的改变：'
)

_bullet(raw_doc, '提升培训任务执行率。通过课程部署、完成度追踪与积分约束，让每一项培训任务都有'
        '明确责任人和截止时间，结果有迹可查。')
_bullet(raw_doc, '提升员工学习参与度。引入正负积分机制，以数据化方式建立轻量级激励和行为约束，'
        '代替纯粹的行政通知。')
_bullet(raw_doc, '支持视频化培训交付。企业可上传操作规程、安全演示等视频内容，系统记录每位员工的'
        '观看进度与完成状态，全程留痕。')
_bullet(raw_doc, '实现管理可视化。培训矩阵将员工、岗位、职能与课程完成情况整合为一张实时视图，'
        '管理者无需再依赖手工 Excel 统计。')
_bullet(raw_doc, '降低无效学习风险。通过轻量化防挂机机制，减少"打开课程后离开""视频拖到结尾"'
        '等形式化学习行为，提升培训数据的真实性。')
_bullet(raw_doc, '预留合规扩展空间。本期建立的培训记录、考试台账和积分流水，可在后续阶段直接扩展'
        '为迎检台账、资质档案和 ESG 培训数据出口。')

_add_key_callout(doc,
    '本期建设的本质是构建一条完整的培训管理链路：内容可管理 → 任务可部署 → 过程可监督 '
    '→ 结果可统计 → 数据可导出。',
    label='建设主线  ')

# ── 二、本期建设范围 ──────────────────────────────────────────────────────────
doc.add_heading('二、本期建设范围', level=1)
doc.add_body(
    '本期建设围绕七个功能版块展开。其中培训内容管理与课程部署是业务运转的基础；积分体系、视频培训、'
    '培训矩阵和防挂机机制是本期核心增强点；数据统计与台账输出则是管理价值落地的最终出口。'
)

doc.add_data_table(
    headers=['模块', '本期建设内容', '对 Vesuvius 的业务价值'],
    data_rows=[
        ['培训内容管理',
         '课程内容编辑、完成规则配置、内容审核流、版本记录',
         '保证课程内容可控发布，降低版本混乱和未授权修改风险'],
        ['课程部署',
         '按人员/部门/岗位/职能分配课程，设置截止时间与提醒规则',
         '将课程从静态资源转化为有责任人、有时间节点的执行任务'],
        ['正负积分体系',
         '完成加分、逾期扣分、积分流水查询、管理端统计',
         '以数据化激励代替重复行政催促，形成正向学习文化'],
        ['视频上传与播放',
         '视频上传、在线播放、断点续播、观看进度记录与完成规则',
         '支持操作规程、安全培训等视频化内容，过程留痕满足合规需要'],
        ['培训矩阵',
         '按员工/部门/岗位/职能展示课程完成率，支持筛选与 Excel 导出',
         '一张矩阵替代多份 Excel，管理者实时掌握培训覆盖缺口'],
        ['学习过程防挂机',
         '无操作提醒、页面切出记录、视频进度限制、过程复核记录',
         '提升学习真实性，确保培训记录的可信度与可审计性'],
        ['数据统计与台账',
         '完成报表、考试记录、积分明细、视频学习记录、Excel 导出',
         '为内部管理提供数据支撑，为后续迎检台账扩展打好基础'],
    ],
    col_widths_mm=[28, 65, 62],
)

doc.add_body('')
doc.add_note(
    '人脸识别、活体检测、多清晰度转码、积分商城等能力暂不纳入本期默认范围，'
    '可在后续阶段按需评估。一期聚焦核心闭环，避免过度建设带来的维护成本。',
    label='一期边界说明：',
)

# ── 三、功能方案详解 ──────────────────────────────────────────────────────────
doc.add_heading('三、功能方案详解', level=1)

# 3.1 培训内容管理
doc.add_heading('3.1  培训内容管理', level=2)
doc.add_body(
    '培训内容管理是平台的内容基础设施。方案的重点不在于堆砌编辑器功能，而在于建立可控的内容发布'
    '流程——确保每一门课程都经过审核、有版本记录、且完成规则清晰。'
)
doc.add_body(
    '完成规则建议采用"组合配置"模式：制度宣导类课程可设为"阅读确认即完成"；安全培训类课程'
    '可叠加"视频观看比例 + 考试通过"；资质认证类课程可进一步要求"证书提交或复训记录"。'
    '这种弹性设计让平台真正适配 Vesuvius 多样化的培训场景，而不是强迫所有课程走同一套标准。'
)
doc.add_data_table(
    headers=['功能', '一期建设说明'],
    data_rows=[
        ['课程内容编辑',   '支持标题、简介、章节、页面内容、图片、附件与视频的编辑和维护'],
        ['课程完成设置',   '可组合配置：必修章节 / 视频观看比例 / 考试通过 / 最低学习时长等'],
        ['考试与测验',     '支持章节小测与课程考试；预留 AI 辅助出题接口，后续可一键生成题库'],
        ['内容审核流',     '课程发布前经培训负责人审核，支持通过、退回与发布记录'],
        ['版本管理',       '每次发布形成版本记录（版本号、发布时间、发布人），历史版本可追溯'],
        ['附件管理',       '每门课程可关联文档、图片等参考资料，建议默认支持 10 份附件/课程'],
    ],
    col_widths_mm=[38, 117],
)

# 3.2 课程部署
doc.add_heading('3.2  课程部署', level=2)
doc.add_body(
    '课程部署是连接"内容资源"与"学习任务"的关键环节。在没有部署模块的情况下，课程发布后'
    '管理员只能靠人工通知、微信群催促；有了部署模块，每次培训都能明确"谁必须学、什么时候学完、'
    '没学完怎么提醒"，执行责任从人工跟进转移到系统自动化管理。'
)
doc.add_data_table(
    headers=['功能', '一期建设说明'],
    data_rows=[
        ['部署对象',     '支持按员工、部门、岗位、职能选择培训对象，可根据客户组织架构扩展'],
        ['时间与截止',   '设置开始时间、截止时间，支持立即生效或定时开始'],
        ['完成规则绑定', '引用课程级完成规则，并可在任务级补充截止约束'],
        ['学习提醒',     '针对未开始、临近截止、逾期未完成状态发送站内提醒；企业微信/短信可扩展'],
        ['完成跟踪',     '管理员可实时查看完成率、未完成人员清单、逾期情况与考试通过状态'],
        ['部署记录',     '每次部署保留对象快照、课程版本、时间配置与完成情况，支持导出'],
    ],
    col_widths_mm=[38, 117],
)

# 3.3 积分体系
doc.add_heading('3.3  正负积分体系', level=2)
doc.add_body(
    '积分体系的设计理念是"轻量激励，不做复杂系统"。一期目标不是建设积分商城或绩效联动，'
    '而是先把学习行为的数字化记录做扎实——完成了什么、什么时候完成、因为什么被加减分，'
    '每一条都有流水可查，让管理者和员工都看得见。'
)
doc.add_data_table(
    headers=['积分类型', '规则示例', '处理方式'],
    data_rows=[
        ['正向积分', '首次登录 +5；完成课程 +10；考试通过 +10；按时完成 +5', '系统按学习行为自动写入积分流水'],
        ['负向积分', '课程逾期 -5；考试未通过 -3；超截止时间 -5',           '按部署截止时间和考试结果自动触发'],
        ['人工调整', '活动奖励、补发、纠错扣减',                           '保留操作人、原因与时间，全程可审计'],
        ['员工视图', '个人积分总量、积分明细、近期变化趋势',                 '学员端可见，增强透明度与参与感'],
        ['管理视图', '部门积分、员工排行、积分来源分布',                     '支持查询、筛选与 Excel 导出'],
    ],
    col_widths_mm=[28, 72, 55],
)
_add_key_callout(doc,
    '后续如需建设积分商城、班组 PK 或"年度学习之星"评选，本期积分流水数据可直接作为基础。'
    '一期的克制是为了后期的可扩展。',
    label='扩展路径  ')

# 3.4 视频上传与播放
doc.add_heading('3.4  视频上传与播放', level=2)
doc.add_body(
    '视频培训在 Vesuvius 的操作规程培训、安全演示和设备维护场景中具有重要价值。本期建设'
    '优先确保视频培训的核心体验稳定可用：上传流畅、播放可靠、进度记录准确。'
)
doc.add_data_table(
    headers=['功能', '一期建设说明'],
    data_rows=[
        ['视频上传',     '管理端在课程章节中上传视频，优先支持 MP4 / H.264 格式'],
        ['对象存储',     '视频文件接入对象存储，与业务服务器分离，便于后续扩容'],
        ['在线播放',     '支持 PC 与移动端浏览器播放，无需安装插件'],
        ['断点续播',     '记录上次观看位置，学员下次打开自动从断点继续'],
        ['观看进度记录', '实时记录观看比例与时长，用于课程完成判断'],
        ['播放控制',     '可配置是否允许拖拽、是否强制观看到指定比例（如 80%）'],
    ],
    col_widths_mm=[38, 117],
)
doc.add_body('')
doc.add_note(
    '若对高并发播放、多清晰度切换或 CDN 分发有较高要求，建议评估视频云服务方案（如阿里云点播）。'
    '该部分会影响建设成本与后续运维投入，需结合贵司实际使用规模决策。',
    label='特殊场景说明：',
)

# 3.5 培训矩阵
doc.add_heading('3.5  培训矩阵', level=2)
doc.add_body(
    '培训矩阵要解决的不是数据收集问题——那些数据已经在系统里了。它要解决的是"管理者每次'
    '想了解培训覆盖情况，都要从头拉数据、手动处理 Excel"的问题。培训矩阵是那张一直存在、'
    '随时可看、无需人工整理的实时视图。'
)
doc.add_data_table(
    headers=['矩阵维度', '展示内容', '典型用途'],
    data_rows=[
        ['员工维度', '姓名、部门、岗位、职能、已完成课程数、完成率、未完成课程',     '查看个人培训进度与缺口'],
        ['部门维度', '部门人数、应训人数、完成率、逾期人数',                       '部门负责人掌握整体执行情况'],
        ['岗位维度', '岗位应训课程、完成人数、未完成人数',                         '评估岗位培训覆盖是否达标'],
        ['课程维度', '部署批次、完成率、考试通过率、逾期人数',                     '评估特定课程的执行效果'],
        ['状态维度', '未开始 / 学习中 / 已完成 / 逾期未完成 / 考试未通过',         '快速筛选定位待处理人员'],
    ],
    col_widths_mm=[28, 75, 52],
)

# 3.6 学习过程防挂机
doc.add_heading('3.6  学习过程防挂机机制', level=2)
doc.add_body(
    '防挂机机制的出发点不是监控员工，而是保障培训记录的可信度。一旦培训记录被用于合规审查'
    '或认证申请，"人在系统外、数据显示已完成"的情况会让整套数据失去可信度。本期采用轻量化'
    '方案，在不涉及摄像头权限和员工隐私的前提下，实现基础层面的学习真实性验证。'
)
doc.add_data_table(
    headers=['机制', '触发条件', '系统处理'],
    data_rows=[
        ['长时间无操作提醒',   '学习页面持续无鼠标 / 键盘 / 触控操作',         '弹出确认提示；未确认则暂停有效学习计时'],
        ['页面切出提示',       '学习中切换标签页或最小化浏览器',               '记录次数；超限时提示返回或暂停进度'],
        ['视频拖拽限制',       '学员跳过未观看的视频片段',                     '不计入已观看进度；完成判断基于真实观看比例'],
        ['弹窗随机确认',       '学习过程中定时或随机触发',                     '未确认则暂停课程进度累计'],
        ['过程记录',           '学习全程',                                    '管理端可查看学习时长、切出次数、确认情况'],
    ],
    col_widths_mm=[38, 52, 65],
)
doc.add_body('')
doc.add_note(
    '人脸识别、活体检测等强监管能力涉及员工隐私与摄像头权限，适合在贵司明确提出强合规要求后'
    '作为专项能力评估，本期暂不纳入默认范围。',
    label='隐私合规说明：',
)

# 3.7 数据统计与台账
doc.add_heading('3.7  数据统计与台账输出', level=2)
doc.add_body(
    '数据输出是培训平台价值的最终证明。无论是每月内部汇报、年度安全培训统计，还是迎接政府'
    '安全检查或 ISO 审核，系统能否快速提供一份清晰完整的台账，决定了管理团队在关键时刻的'
    '响应效率。'
)
doc.add_data_table(
    headers=['报表类型', '主要字段', '应用场景'],
    data_rows=[
        ['培训完成报表',   '员工、部门、岗位、课程、完成状态、完成时间、完成率',     '日常培训执行管理'],
        ['考试结果报表',   '考试名称、分数、是否通过、考试次数、提交时间',           '学习效果核查与补训决策'],
        ['积分明细报表',   '员工、积分变化、来源、课程、操作人',                   '积分追溯与管理复核'],
        ['视频学习记录',   '视频名称、观看比例、观看时长、是否达标',               '视频课程完成判断依据'],
        ['部署任务报表',   '课程版本、部署对象、开始 / 截止时间、完成情况',         '培训任务执行复盘'],
        ['台账基础导出',   '按人员 / 课程 / 部门 / 时间段多维度导出',             '迎检台账基础数据 / ESG 培训数据'],
    ],
    col_widths_mm=[32, 72, 51],
)

# ── 四、技术实现方案 ──────────────────────────────────────────────────────────
doc.add_heading('四、技术实现方案', level=1)
doc.add_body(
    '本项目采用前后端分离架构，围绕课程、部署、学习记录、积分、视频资源和报表统计建立清晰的'
    '数据模型。技术选型优先保证可维护性、可扩展性与可审计性，避免为追求技术新颖而增加不必要'
    '的复杂度。'
)

doc.add_heading('4.1  前端架构', level=3)
_bullet(raw_doc, '管理端：课程内容编辑、审核发布、课程部署、积分配置、培训矩阵、报表导出。')
_bullet(raw_doc, '学员端：课程学习、视频播放、章节测验、考试、积分查看、学习任务清单。')
_bullet(raw_doc, '响应式支持：优先保障 PC 端体验；移动端浏览器访问按贵司实际使用场景确认适配范围。')

doc.add_heading('4.2  后端服务', level=3)
_bullet(raw_doc, '课程服务：管理课程内容、章节结构、完成规则、审核状态与版本历史。')
_bullet(raw_doc, '部署服务：管理培训任务、分配对象、时间规则、提醒触发与完成统计。')
_bullet(raw_doc, '学习记录服务：记录课程进度、视频观看、测验考试与过程行为信息。')
_bullet(raw_doc, '积分服务：根据学习行为自动生成积分流水，支持人工调整与完整审计日志。')
_bullet(raw_doc, '报表服务：按员工、部门、岗位、课程、时间维度聚合统计并支持导出。')

doc.add_heading('4.3  文件与视频存储', level=3)
_bullet(raw_doc, '文档与图片统一接入文件资源服务，支持权限控制与后续扩容。')
_bullet(raw_doc, '视频文件采用对象存储方案，与业务服务器磁盘完全分离。')
_bullet(raw_doc, '基础版支持 MP4 原文件播放；如需 HLS 分片、转码或 CDN 加速，可在增强阶段按需接入。')

doc.add_heading('4.4  权限与数据安全', level=3)
_bullet(raw_doc, '角色体系：管理员、培训负责人、审核人、部门负责人、普通学员各具不同操作权限。')
_bullet(raw_doc, '关键操作审计：课程发布、积分人工调整、部署撤销等操作均记录操作人与时间戳。')
_bullet(raw_doc, '数据保护：学习记录、考试成绩与积分流水不可随意修改；必要修改需保留原因与操作人，'
        '满足 GDPR 相关的数据可追溯要求。')
_bullet(raw_doc, '私有化部署：系统数据存储在贵司环境内，不向第三方上传原始员工数据。')

# ── 五、我们的交付承诺 ────────────────────────────────────────────────────────
doc.add_heading('五、我们的交付承诺', level=1)
doc.add_body(
    '选择一个技术伙伴，考量的不只是功能清单，更是对方能否在约定时间内稳定交付、遇到问题时'
    '是否及时响应、项目结束后是否还能依赖。以下是我们的具体承诺。'
)

doc.add_heading('5.1  结构化交付流程', level=3)
_bullet(raw_doc, '项目启动前完成需求确认会议，形成需求确认清单与功能边界文档，避免后期范围蔓延。')
_bullet(raw_doc, '每周固定同步进度、待确认事项和风险清单，贵司无需主动追问开发状态。')
_bullet(raw_doc, '原型评审阶段提前识别理解偏差，避免开发完成后的返工。')
_bullet(raw_doc, '验收标准在项目启动前与贵司共同确认，不留模糊空间。')

doc.add_heading('5.2  上线后支持', level=3)
_bullet(raw_doc, '上线后 30 天内发现的功能缺陷，免费修复，无额外费用。')
_bullet(raw_doc, '提供操作说明文档与视频导览，降低贵司内部推广成本。')
_bullet(raw_doc, '支持远程培训：管理员使用培训 + 学员端使用演示，确保团队快速上手。')

doc.add_heading('5.3  灵活的扩展路径', level=3)
_bullet(raw_doc, '本期建设的数据模型和接口设计充分考虑后续扩展，避免二期需要重写基础架构。')
_bullet(raw_doc, '如后续需要对接贵司现有 HRMS（SAP、Workday 等）同步组织架构，接口预留成本极低。')
_bullet(raw_doc, '积分商城、岗位风险匹配、AI 课件生成等增强能力均可在现有基础上平滑扩展。')

_add_key_callout(doc,
    '我们建立的不是单次项目交付关系，而是一个随贵司业务成长可持续迭代的数字培训基础设施。',
    label='合作理念  ')

# ── 六、实施计划 ──────────────────────────────────────────────────────────────
doc.add_heading('六、实施计划', level=1)
doc.add_body(
    '按本期标准范围估算，项目建议实施周期为 8 周。若最终范围有所调整（如增加视频云服务、'
    'HRMS 数据对接或历史数据迁移），周期与报价将相应评估。'
)

doc.add_data_table(
    headers=['阶段', '周次', '主要工作', '交付物'],
    data_rows=[
        ['需求确认',
         '第 1 周',
         '确认功能范围、角色权限、组织数据结构、报价版本与验收口径',
         '需求确认清单、项目计划'],
        ['原型与设计',
         '第 2 周',
         '完成主要页面原型、数据流与接口范围设计',
         '页面原型、功能说明、接口清单'],
        ['核心功能开发',
         '第 3–6 周',
         '开发七大功能模块：内容管理、课程部署、积分体系、视频播放、培训矩阵、防挂机、报表',
         '可联调功能版本'],
        ['联调与测试',
         '第 7 周',
         '前后端联调、功能测试、数据测试、权限测试、导出测试',
         '测试记录、问题修复清单'],
        ['试运行与验收',
         '第 8 周',
         '客户试用、问题修复、验收确认、使用培训',
         '验收版本、操作说明、交付清单'],
    ],
    col_widths_mm=[28, 16, 72, 39],
)
doc.add_body('')
doc.add_note(
    '项目实施期间，需贵司指定业务负责人参与需求确认、原型评审、测试验收和上线试运行，'
    '预计每周沟通时间约 1–2 小时。高效的客户配合是项目按时交付的重要保障。',
    label='客户侧配合：',
)

# ── 七、下一步行动 ────────────────────────────────────────────────────────────
doc.add_heading('七、下一步行动', level=1)
doc.add_body(
    '如贵司对本建议书的方向认可，建议通过以下三步推进项目落地：'
)

doc.add_data_table(
    headers=['步骤', '动作', '建议时间', '贵司对接人'],
    data_rows=[
        ['Step 1', '安排范围确认会议，逐项确认功能细节与技术边界', '建议 1 周内', 'HR / 培训负责人'],
        ['Step 2', '基于确认范围出具正式报价单，双方签署合作协议', '建议 2 周内', 'HR / 采购'],
        ['Step 3', '项目正式启动，我方提交项目计划与第一版原型',    '签约后 1 周', 'IT / 业务负责人'],
    ],
    col_widths_mm=[17, 72, 30, 36],
)

doc.add_body('')
_add_key_callout(doc,
    '我们已为 Vesuvius 中国区预备了专项资源。如需提前了解平台演示或进行技术对话，'
    '欢迎随时安排。期待与贵团队建立长期合作。',
    label='联系我们  ')

doc.add_spacer(2)
doc.add_signature_block(company=BT.BRAND_FULL_CN, show_seal=True)

# ── 保存 ──────────────────────────────────────────────────────────────────────
OUTPUT = os.path.join(_DOCS, 'Vesuvius_培训平台增强建议书_v1.0.docx')
doc.save(OUTPUT)
print(f'✓ Saved: {OUTPUT}')
