"""
gen_proposal_final_docx.py
Generates: Vesuvius_培训平台增强建议书_v3_final.docx

Differences from v3.0:
  - All data tables use "rounded-card" style:
      · Outer border: thick brand-green (#3EC99E, sz=10) — card frame
      · No inner vertical dividers (insideV = none) — modern, open feel
      · Thin light-gray horizontal row separators (insideH, sz=4)
      · Generous cell padding (top/bottom 100 twips, left/right 120 twips)
  - Table headers simulate a primary-color gradient:
      · Fill: solid primary green (#3EC99E)
      · Top border: thick lighter green (#8FDFC5, sz=20) — top highlight
      · Together they read as a gradient from light → solid
  - Same content & pricing as v3.0  (Year 1: ¥178,000 / Year 2+: ¥140,000)
"""

import sys, os, io

_BRAND   = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
_PROJECT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, _BRAND)

_DOCS  = os.path.join(_PROJECT, 'docs')
_MEDIA = os.path.join(_PROJECT, 'media')

os.makedirs(_DOCS, exist_ok=True)

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from scripts.docx_builder import (
    BrandDocx, _set_font, _cell_shading, _cell_borders,
    _para_left_bar, _para_shading,
    _GRAY, _GREEN, _NONE,
)
import scripts.brand_tokens as BT

_CN_FONT = 'Alibaba PuHuiTi 2.0'
_EN_FONT = 'Inter'

# Gradient highlight colour for header top border (lighter than primary green)
_HDR_HIGHLIGHT = '8FDFC5'   # ~60% tint of #3EC99E


# ── Font helpers ───────────────────────────────────────────────────────────────

def _set_run_font(run, font_cn=_CN_FONT, font_en=_EN_FONT):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    font_en)
    rFonts.set(qn('w:hAnsi'),    font_en)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:cs'),       font_cn)
    run.font.name = font_en


def _set_doc_font(raw_doc):
    styles_elem = raw_doc.styles._element
    docDefs = styles_elem.find(qn('w:docDefaults'))
    if docDefs is not None:
        rPrDef = docDefs.find(qn('w:rPrDefault'))
        if rPrDef is None:
            rPrDef = OxmlElement('w:rPrDefault')
            docDefs.append(rPrDef)
        rPr = rPrDef.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDef.append(rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'),    _EN_FONT)
        rFonts.set(qn('w:hAnsi'),    _EN_FONT)
        rFonts.set(qn('w:eastAsia'), _CN_FONT)
        rFonts.set(qn('w:cs'),       _CN_FONT)
    try:
        normal = raw_doc.styles['Normal']
        normal.font.name = _EN_FONT
        rPr2 = normal.font._element
        rFonts2 = rPr2.find(qn('w:rFonts'))
        if rFonts2 is None:
            rFonts2 = OxmlElement('w:rFonts')
            rPr2.insert(0, rFonts2)
        rFonts2.set(qn('w:eastAsia'), _CN_FONT)
    except Exception:
        pass
    for para in raw_doc.paragraphs:
        for run in para.runs:
            _set_run_font(run)
    for tbl in raw_doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run)


def _inject_saved_header(docx_path, src_docx_path):
    import zipfile, shutil
    with zipfile.ZipFile(src_docx_path, 'r') as src_z:
        hdr_bytes = src_z.read('word/header1.xml')
        hdr_rels  = src_z.read('word/_rels/header1.xml.rels')
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == 'word/header1.xml':
                zout.writestr(item, hdr_bytes)
            elif item.filename == 'word/_rels/header1.xml.rels':
                zout.writestr(item, hdr_rels)
            else:
                zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp, docx_path)


# ── Content helpers ────────────────────────────────────────────────────────────

def _bullet(raw_doc, text, level=0, sz=10.5):
    p = raw_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    indent = Pt(16 + level * 14)
    p.paragraph_format.left_indent       = indent
    p.paragraph_format.first_line_indent = Pt(-13)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.35
    bullet_char = '•' if level == 0 else '–'
    r1 = p.add_run(f'{bullet_char}  ')
    _set_font(r1, size_pt=sz, color_hex=BT.PRIMARY_500_HEX, bold=True)
    _set_run_font(r1)
    r2 = p.add_run(text)
    _set_font(r2, size_pt=sz, color_hex=BT.NEUTRAL_700_HEX)
    _set_run_font(r2)
    return p


def _page_break(raw_doc):
    p = raw_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════════════════
# Rounded-card table helpers
# ══════════════════════════════════════════════════════════════════════════════

def _get_or_add_tblPr(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    return tblPr


def _apply_card_style(tbl):
    """
    Apply rounded-card look to a table:
      - Outer border: thick brand-green
      - insideV: none  (no column dividers)
      - insideH: thin light-gray  (row separators)
      - Cell margin: generous
    Call AFTER creating the table; override any existing tblBorders.
    """
    PRIMARY = BT.PRIMARY_500_HEX.lstrip('#')
    GRAY    = BT.NEUTRAL_200_HEX.lstrip('#')

    tblPr = _get_or_add_tblPr(tbl)

    # Remove any existing tblBorders so we start clean
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)

    tblBdr = OxmlElement('w:tblBorders')
    # Outer borders: thick green
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '10')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), PRIMARY)
        tblBdr.append(el)
    # insideV: none — no column dividers
    iv = OxmlElement('w:insideV')
    iv.set(qn('w:val'),   'none')
    iv.set(qn('w:sz'),    '0')
    iv.set(qn('w:space'), '0')
    iv.set(qn('w:color'), 'FFFFFF')
    tblBdr.append(iv)
    # insideH: thin gray — row separators
    ih = OxmlElement('w:insideH')
    ih.set(qn('w:val'),   'single')
    ih.set(qn('w:sz'),    '4')
    ih.set(qn('w:space'), '0')
    ih.set(qn('w:color'), GRAY)
    tblBdr.append(ih)
    tblPr.append(tblBdr)

    # Generous cell margins (twips)
    old_mar = tblPr.find(qn('w:tblCellMar'))
    if old_mar is not None:
        tblPr.remove(old_mar)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', 100), ('bottom', 100), ('left', 120), ('right', 120)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


def _gradient_header_cell(cell):
    """
    Simulate a gradient on a header cell:
      fill = primary green (#3EC99E)
      top border = lighter green (#8FDFC5, thick) — highlights the top edge
    Together they read as a gradient from light-green (top) → solid green (body).
    """
    _cell_shading(cell, BT.PRIMARY_500_HEX)

    tcPr = cell._tc.get_or_add_tcPr()
    # Remove any existing tcBorders before adding ours
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcBdr = OxmlElement('w:tcBorders')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'),   'single')
    top_el.set(qn('w:sz'),    '20')
    top_el.set(qn('w:space'), '0')
    top_el.set(qn('w:color'), _HDR_HIGHLIGHT)
    tcBdr.append(top_el)
    tcPr.append(tcBdr)


def _fancy_table(raw_doc, headers, data_rows, col_widths_mm=None):
    """
    Rounded-card data table with gradient-style header row.
    Replaces doc.add_data_table() for this document.
    """
    n_cols = len(headers)
    tbl = raw_doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Use Table Normal (no built-in borders) so we control everything
    try:
        tbl.style = raw_doc.styles['Table Normal']
    except KeyError:
        pass  # fallback: python-docx default

    _apply_card_style(tbl)

    # ── Header row ────────────────────────────────────────────────────────────
    hrow = tbl.rows[0]
    for j, hdr_text in enumerate(headers):
        c = hrow.cells[j]
        if col_widths_mm:
            c.width = Mm(col_widths_mm[j])
        _gradient_header_cell(c)
        p = c.paragraphs[0]
        p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(hdr_text)
        _set_font(r, size_pt=10, bold=True, color_hex=BT.TABLE_HEADER_FG)
        _set_run_font(r)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for i, row_data in enumerate(data_rows):
        tr = tbl.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            c = tr.cells[j]
            if col_widths_mm:
                c.width = Mm(col_widths_mm[j])
            if i % 2 == 1:
                _cell_shading(c, BT.TABLE_STRIPE_BG)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(cell_text)
            _set_font(r, size_pt=10.5, color_hex=BT.NEUTRAL_700_HEX)
            _set_run_font(r)

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# Cover page  （参考 URS 咨询报告风格）
# ══════════════════════════════════════════════════════════════════════════════

def _cover(raw_doc, vesuvius_logo_path, our_logo_png_path):

    header_tbl = raw_doc.add_table(1, 2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in header_tbl.rows[0].cells:
        _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

    lc_h = header_tbl.rows[0].cells[0]
    lc_h.width = Mm(100)
    lp_h = lc_h.paragraphs[0]
    lp_h.paragraph_format.space_before = Pt(0)
    lp_h.paragraph_format.space_after  = Pt(0)
    r_type = lp_h.add_run('S O L U T I O N   P R O P O S A L')
    _set_font(r_type, size_pt=8, color_hex=BT.NEUTRAL_400_HEX)
    _set_run_font(r_type)

    rc_h = header_tbl.rows[0].cells[1]
    rc_h.width = Mm(55)
    rp_h = rc_h.paragraphs[0]
    rp_h.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
    rp_h.paragraph_format.space_before = Pt(0)
    rp_h.paragraph_format.space_after  = Pt(0)
    r_conf = rp_h.add_run('机密文件，仅供内部参阅')
    _set_font(r_conf, size_pt=8, color_hex=BT.NEUTRAL_400_HEX, italic=True)
    _set_run_font(r_conf)

    rule1 = raw_doc.add_paragraph()
    rule1.paragraph_format.space_before = Pt(10)
    rule1.paragraph_format.space_after  = Pt(0)
    pPr = rule1._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BT.PRIMARY_500_HEX.lstrip('#'))
    bdr.append(bot)
    pPr.append(bdr)

    for _ in range(2):
        sp = raw_doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    if os.path.exists(vesuvius_logo_path):
        lp = raw_doc.add_paragraph()
        lp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(20)
        lp.add_run().add_picture(vesuvius_logo_path, width=Mm(58))

    title_p = raw_doc.add_paragraph()
    title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(6)
    r_title = title_p.add_run('企业培训平台功能增强建设建议书')
    _set_font(r_title, size_pt=26, bold=True, color_hex=BT.NEUTRAL_900_HEX)
    _set_run_font(r_title)

    en_p = raw_doc.add_paragraph()
    en_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    en_p.paragraph_format.space_before = Pt(0)
    en_p.paragraph_format.space_after  = Pt(4)
    r_en = en_p.add_run('Enterprise Training Platform Enhancement Proposal')
    _set_font(r_en, size_pt=11, color_hex=BT.NEUTRAL_400_HEX)
    _set_run_font(r_en)

    scope_p = raw_doc.add_paragraph()
    scope_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
    scope_p.paragraph_format.space_before = Pt(2)
    scope_p.paragraph_format.space_after  = Pt(0)
    _para_shading(scope_p, BT.NEUTRAL_100_HEX)
    r_scope = scope_p.add_run(
        '   培训内容管理  ·  课程部署  ·  积分激励  ·  视频培训  ·  培训矩阵  ·  数据台账'
    )
    _set_font(r_scope, size_pt=9.5, color_hex=BT.NEUTRAL_700_HEX)
    _set_run_font(r_scope)

    for _ in range(5):
        raw_doc.add_paragraph().paragraph_format.space_after = Pt(0)

    addr_tbl = raw_doc.add_table(1, 2)
    addr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c in addr_tbl.rows[0].cells:
        _cell_borders(c, top=_GRAY, bottom=_NONE, left=_NONE, right=_NONE)

    lc = addr_tbl.rows[0].cells[0]
    lc.width = Mm(80)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    lp_label = lc.paragraphs[0]
    lp_label.paragraph_format.space_before = Pt(10)
    lp_label.paragraph_format.space_after  = Pt(4)
    r_to = lp_label.add_run('提交至')
    _set_font(r_to, size_pt=8, color_hex=BT.NEUTRAL_400_HEX, italic=True)
    _set_run_font(r_to)

    def _addr_line(cell, text, bold=False, sz=10):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        _set_font(r, size_pt=sz, bold=bold,
                  color_hex=BT.NEUTRAL_900_HEX if bold else BT.NEUTRAL_700_HEX)
        _set_run_font(r)

    _addr_line(lc, '维苏威高级陶瓷（中国）有限公司', bold=True, sz=10.5)
    _addr_line(lc, '（总部及核心办公室）', sz=9.5)
    _addr_line(lc, '江苏省苏州市工业园区星明街 221 号', sz=9.5)

    rc = addr_tbl.rows[0].cells[1]
    rc.width = Mm(75)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    rp_label = rc.paragraphs[0]
    rp_label.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
    rp_label.paragraph_format.space_before = Pt(10)
    rp_label.paragraph_format.space_after  = Pt(4)
    r_by = rp_label.add_run('提交方')
    _set_font(r_by, size_pt=8, color_hex=BT.NEUTRAL_400_HEX, italic=True)
    _set_run_font(r_by)

    if os.path.exists(our_logo_png_path):
        logo_p = rc.add_paragraph()
        logo_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after  = Pt(4)
        logo_p.add_run().add_picture(our_logo_png_path, width=Mm(32))

    def _addr_line_right(cell, text, bold=False, sz=9.5):
        p = cell.add_paragraph()
        p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        _set_font(r, size_pt=sz, bold=bold,
                  color_hex=BT.NEUTRAL_900_HEX if bold else BT.NEUTRAL_700_HEX)
        _set_run_font(r)

    _addr_line_right(rc, BT.BRAND_FULL_CN, bold=True, sz=10)
    _addr_line_right(rc, '上海漕河泾创业中心大厦（桂平路）302 栋', sz=9.5)

    date_p = raw_doc.add_paragraph()
    date_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_before = Pt(14)
    date_p.paragraph_format.space_after  = Pt(0)
    r_date = date_p.add_run('2026 年 6 月')
    _set_font(r_date, size_pt=9, color_hex=BT.NEUTRAL_400_HEX)
    _set_run_font(r_date)

    _page_break(raw_doc)


# ══════════════════════════════════════════════════════════════════════════════
# 目录页
# ══════════════════════════════════════════════════════════════════════════════

def _toc(raw_doc):
    toc_title = raw_doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(0)
    toc_title.paragraph_format.space_after  = Pt(16)
    r_tt = toc_title.add_run('目  录')
    _set_font(r_tt, size_pt=20, bold=True, color_hex=BT.NEUTRAL_900_HEX)
    _set_run_font(r_tt)

    rule = raw_doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(20)
    pPr = rule._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BT.PRIMARY_500_HEX.lstrip('#'))
    bdr.append(bot)
    pPr.append(bdr)

    toc_data = [
        ('一',  '项目背景与目标',
         '建设主线与六大管理改变',                      False),
        ('二',  '本期建设范围',
         '七个功能版块总览',                            False),
        ('三',  '功能方案',
         '各模块功能详述与价值说明',                    False),
        ('',    '3.1  培训内容管理',
         '完成规则 · 审核流 · 版本记录 · AI 辅助',      True),
        ('',    '3.2  课程部署',
         '任务分配 · 截止追踪 · 部署历史',              True),
        ('',    '3.3  正负积分体系',
         '激励约束 · 积分流水 · 团队参与度',            True),
        ('',    '3.4  视频培训',
         '在线播放 · 断点续播 · 进度真实记录',          True),
        ('',    '3.5  培训矩阵',
         '全员覆盖视图 · 多维度筛选 · 一键导出',        True),
        ('',    '3.6  数据统计与台账',
         '多类型报表 · 合规数据出口',                   True),
        ('四',  '双方团队分工',
         '任务清单 · 工作界面 · 初步项目进度',          False),
        ('五',  '报价方案',
         '分层席位模型 · 费用明细 · 报价说明',          False),
    ]

    toc_tbl = raw_doc.add_table(len(toc_data), 3)
    toc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (num, title, desc, is_sub) in enumerate(toc_data):
        row = toc_tbl.rows[i]
        for c in row.cells:
            _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

        c0 = row.cells[0]
        c0.width = Mm(12)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0 if is_sub else 6)
        p0.paragraph_format.space_after  = Pt(2)
        if num:
            r0 = p0.add_run(num)
            _set_font(r0, size_pt=11, bold=True, color_hex=BT.PRIMARY_500_HEX)
            _set_run_font(r0)

        c1 = row.cells[1]
        c1.width = Mm(70)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0 if is_sub else 6)
        p1.paragraph_format.space_after  = Pt(2)
        indent = Pt(6) if is_sub else Pt(0)
        p1.paragraph_format.left_indent  = indent
        r1 = p1.add_run(title)
        sz = 10 if is_sub else 11
        col = BT.NEUTRAL_700_HEX if is_sub else BT.NEUTRAL_900_HEX
        _set_font(r1, size_pt=sz, bold=(not is_sub), color_hex=col)
        _set_run_font(r1)

        c2 = row.cells[2]
        c2.width = Mm(73)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(0 if is_sub else 6)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(desc)
        _set_font(r2, size_pt=9, color_hex=BT.NEUTRAL_400_HEX)
        _set_run_font(r2)

        if not is_sub and num:
            for c in row.cells:
                _cell_shading(c, BT.NEUTRAL_100_HEX)

    _page_break(raw_doc)


# ══════════════════════════════════════════════════════════════════════════════
# Build document
# ══════════════════════════════════════════════════════════════════════════════

doc     = BrandDocx(doc_type='解决方案建议书')
raw_doc = doc.document

sec = raw_doc.sections[0]
sec.different_first_page_header_footer = True

for tbl in sec.footer.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text and ('人力资源' in run.text or '证明使用' in run.text):
                        run.text = '机密文件 · 仅供 Vesuvius 中国区内部参阅'

_VESUVIUS_LOGO = os.path.join(_MEDIA, 'covers', 'VESUVIUS_logo.png')
_OUR_LOGO      = BT.LOGO_HORIZONTAL_PRIMARY_PNG

# ── 封面 ──────────────────────────────────────────────────────────────────────
_cover(raw_doc, _VESUVIUS_LOGO, _OUR_LOGO)

# ── 目录 ──────────────────────────────────────────────────────────────────────
_toc(raw_doc)

# ── 正文标题信息表 ────────────────────────────────────────────────────────────
doc.add_title(
    '企业培训平台功能增强建设建议书',
    subtitle='ENTERPRISE TRAINING PLATFORM ENHANCEMENT PROPOSAL  ·  FOR VESUVIUS CHINA',
)

doc.add_info_table([
    ('提交至',   '维苏威高级陶瓷（中国）有限公司'),
    ('项目名称', '企业培训平台功能增强建设项目'),
    ('版本',     'V3.0'),
    ('文件性质', '机密文件，仅供内部参阅'),
    ('预计周期', '约 8 周（视最终范围确认）'),
], label_width_mm=35, value_width_mm=110)


# ── 一、项目背景与目标 ─────────────────────────────────────────────────────────
doc.add_heading('一、项目背景与目标', level=1)
doc.add_body(
    '贵司现有培训平台已具备基础的内容管理能力。本次增强的核心目标是让培训从"内容已发布"'
    '真正走向"任务被完成、过程有记录、结果可追溯"——让管理者不再依赖手工 Excel 和'
    '微信群催促，让每一次培训都有闭环。'
)
doc.add_body('本期建设将在六个维度产生直接的管理改变：')

_bullet(raw_doc, '培训任务有人认领、有截止时间、完成情况实时可查，执行率可量化。')
_bullet(raw_doc, '员工参与有正向激励和逾期约束，学习行为有数据记录。')
_bullet(raw_doc, '操作规程、安全演示等视频内容可上传课程并分发部署，观看完成情况自动留痕。')
_bullet(raw_doc, '一张培训矩阵实时展示所有员工的课程覆盖情况，无需人工汇总。')
_bullet(raw_doc, '学习过程真实性有基础保障，培训记录具备可信度与审计能力。')
_bullet(raw_doc, '培训记录、积分流水、考试台账具备向合规档案扩展的数据基础。')

doc.add_note(
    '建设主线：内容可管理 → 任务可部署 → 过程可监督 → 结果可统计 → 数据可导出。',
    label='一句话概括  ',
)

# ── 二、本期建设范围 ──────────────────────────────────────────────────────────
doc.add_heading('二、本期建设范围', level=1)
doc.add_body(
    '本期围绕七个功能版块展开建设，每个版块都直接对应贵司管理团队的一个具体痛点。'
)

_fancy_table(
    raw_doc,
    headers=['功能版块', '对 Vesuvius 的直接价值'],
    data_rows=[
        ['培训内容管理', '课程内容有人审核、有版本记录，确保员工看到的是已授权的最新版本'],
        ['课程部署',     '每次培训都有明确责任人、截止时间和完成追踪，彻底告别微信群催促'],
        ['正负积分体系', '完成课程有奖励、逾期有约束，积分记录透明，驱动员工主动学习'],
        ['视频培训',     '上传操作规程或安全演示视频，员工在线观看，系统自动记录完成状态'],
        ['培训矩阵',     '一张实时视图展示所有员工的课程完成率，部门/岗位缺口一眼可见'],
        ['学习过程监督', '系统记录学习行为真实性，减少挂机和无效观看，保障培训数据可信度'],
        ['数据统计与台账', '完成报表、考试记录、积分明细一键导出，支持内部汇报与合规备查'],
    ],
    col_widths_mm=[38, 117],
)

# ── 三、功能方案 ──────────────────────────────────────────────────────────────
doc.add_heading('三、功能方案', level=1)

doc.add_heading('3.1  培训内容管理', level=2)
doc.add_body(
    '管理员可对课程内容持续维护，设置不同类型课程的完成规则，并通过审核流确保每门课程'
    '在正式发布前经过负责人确认。制度宣导类课程可设为"阅读确认即完成"；安全培训类课程'
    '可叠加"视频观看比例 + 考试通过"；版本记录确保每次改动都有据可查。'
)
_fancy_table(
    raw_doc,
    headers=['功能', '您将获得'],
    data_rows=[
        ['灵活的完成规则', '按课程类型分别配置：阅读确认 / 视频观看比例 / 考试通过，不强迫所有课程走同一套标准'],
        ['内容审核流',     '课程正式发布前须经培训负责人审核，防止未授权内容失误发布到平台内部'],
        ['版本记录',       '每次发布形成版本留痕，历史版本可追溯，迎检时有完整内容变更依据'],
        ['AI 辅助',        '支持 AI 辅助生成知识点新建章节，考试题目支持 AI 辅助生成，大幅降低题库维护工作量'],
    ],
    col_widths_mm=[40, 115],
)

doc.add_heading('3.2  课程部署', level=2)
doc.add_body(
    '课程部署是让培训从"资源库"变成"执行任务"的关键。管理员按部门、岗位或人员维度'
    '分配课程，设置截止时间和提醒规则后，系统自动追踪每个人的完成进度——谁完成了、'
    '谁逾期了、考试结果怎样，全部实时可查，无需手工汇总。'
)
_fancy_table(
    raw_doc,
    headers=['功能', '您将获得'],
    data_rows=[
        ['灵活的分配方式', '支持按员工、部门、岗位、职能、班组多维度分配，一次操作覆盖目标人群'],
        ['截止时间与提醒', '系统在临近截止时自动提醒未完成员工，逾期情况同步通知管理者'],
        ['实时完成追踪',   '管理员随时查看当前部署的完成率、未完成人员清单与考试通过情况'],
        ['部署历史存档',   '每次培训的对象、版本、完成情况长期留存，支持年度培训复盘与审查'],
    ],
    col_widths_mm=[40, 115],
)

doc.add_heading('3.3  正负积分体系', level=2)
doc.add_body(
    '积分体系以数据化方式建立轻量级激励和约束，代替反复的行政催促。完成课程自动加分，'
    '逾期自动扣分，每一笔积分流水都有明细可查——员工看得到自己的表现，管理者看得到'
    '团队的整体参与度。后续如有需要，积分数据可直接扩展为学习排行榜或年度激励评选的依据。'
)
_fancy_table(
    raw_doc,
    headers=['积分规则', '员工视角', '管理者视角'],
    data_rows=[
        ['完成课程 / 考试通过', '个人积分实时增加，可查看积分明细', '整体参与度提升，无需反复催促'],
        ['课程逾期 / 考试未过', '收到系统提示，清楚了解扣分原因', '逾期记录可导出，支持部门绩效参考'],
        ['活动奖励 / 人工调整', '积分变动有通知，透明公正',       '管理端可补发或纠错，全程留审计日志'],
    ],
    col_widths_mm=[45, 60, 50],
)

doc.add_heading('3.4  视频培训', level=2)
doc.add_body(
    '管理员可在课程章节中上传操作规程、安全培训演示或设备维护视频，员工通过 PC 或'
    '手机浏览器在线观看。系统实时记录每位员工的观看进度，可配置"观看达到指定比例'
    '方可完成课程"，确保视频培训留有真实记录而非仅标注已观看。'
)
_fancy_table(
    raw_doc,
    headers=['功能', '您将获得'],
    data_rows=[
        ['在线上传与播放', '视频直接上传至课程，员工无需下载，PC / 手机浏览器均可播放'],
        ['断点续播',       '员工上次观看到哪里，下次自动从断点继续，不影响碎片化学习'],
        ['观看进度记录',   '系统记录实际观看比例，课程完成判断基于真实数据而非点击记录'],
        ['播放控制',       '可配置是否允许拖拽进度条，防止员工跳过关键内容直接标记完成'],
    ],
    col_widths_mm=[40, 115],
)

doc.add_heading('3.5  培训矩阵', level=2)
doc.add_body(
    '培训矩阵解决的是"每次想了解培训覆盖情况，都要从头拉数据、手工处理 Excel"的问题。'
    '它是一张始终在线、实时更新的全员培训视图，管理者可以按部门、岗位或培训状态快速'
    '筛选，一眼定位需要关注的人员或课程，并支持一键导出。'
)
_fancy_table(
    raw_doc,
    headers=['视角', '可以看到什么', '典型场景'],
    data_rows=[
        ['员工个人',   '已完成哪些课程、还有哪些未完成、当前积分',  '员工自主了解学习任务进度'],
        ['部门负责人', '本部门完成率、逾期人数、课程通过率',         '月度培训汇报，无需手工统计'],
        ['培训管理者', '跨部门/岗位的完成缺口、课程执行效果对比',    '安全检查前快速准备培训台账'],
    ],
    col_widths_mm=[28, 82, 45],
)

doc.add_heading('3.6  数据统计与台账', level=2)
doc.add_body(
    '平台将课程完成、考试结果、积分变动、视频观看记录统一汇聚，支持按员工、部门、'
    '岗位或时间段多维度导出。无论是月度内部汇报、年度安全培训统计，还是迎接政府'
    '安全检查或 ISO 审核，都能在几分钟内生成一份完整的数字台账。'
)
_fancy_table(
    raw_doc,
    headers=['报表类型', '典型用途'],
    data_rows=[
        ['培训完成报表', '日常执行管理；按部门/岗位汇总完成率，快速定位缺口'],
        ['考试结果报表', '评估学习效果；识别需要补训的员工'],
        ['积分明细报表', '积分核查与管理复核；支持绩效参考数据导出'],
        ['视频学习记录', '视频课程完成判断依据；满足合规留痕要求'],
        ['培训台账导出', '迎检备查；ESG 培训覆盖数据出口；年度培训统计'],
    ],
    col_widths_mm=[42, 113],
)

# ── 四、双方团队分工 ──────────────────────────────────────────────────────────
doc.add_heading('四、双方团队分工', level=1)

doc.add_heading('4.1  任务清单与工作界面', level=2)
doc.add_body(
    '项目执行期间，双方团队在各阶段均有明确的责任边界。以下任务清单列出了每项工作的分工，'
    '确保推进有序、不产生责任盲区。'
)
_fancy_table(
    raw_doc,
    headers=['编号', '任务描述', '甲方（Vesuvius）职责', '我方职责'],
    data_rows=[
        ['1', '需求确认',
         '确认功能范围与角色权限；指定培训负责人作为项目对接人；审批我方给予的完整需求清单',
         '输出需求清单、主要页面原型及双方认可的验收标准'],
        ['2', '数据与组织准备',
         '提供员工花名册及部门/岗位层级数据；确认课程分类体系与初始管理权限分配',
         '完成组织架构导入、账号初始化及角色权限体系配置'],
        ['3', '功能开发',
         '提供部分课程内容（文档/视频/其他格式）；对原型方案进行业务评审',
         '开发七大功能模块，每周固定进度同步，问题当周响应'],
        ['4', '系统内测',
         '组织 HR 及培训管理员参与内测，提供真实场景操作反馈',
         '根据反馈修复问题并完成回归测试，确保功能符合验收标准'],
        ['5', '上线验收',
         '组织管理员正式验收，确认系统满足建设目标后签署验收确认',
         '提供管理员培训、学员端操作手册及完整交付文件'],
        ['6', '上线后运维',
         '指定系统管理员负责日常运营；统一收集并反馈优化需求',
         '提供上线后 1 年免费维护，功能问题响应时间不超过 1 个工作日'],
    ],
    col_widths_mm=[11, 28, 65, 51],
)

doc.add_heading('4.2  初步项目进度安排', level=2)
doc.add_body('标准版初步建设周期约 8 周，按五个阶段有序推进：')
_fancy_table(
    raw_doc,
    headers=['阶段', '周次', '涉及职能', '里程碑 / 交付物'],
    data_rows=[
        ['需求确认',   '第 1 周',   '产品经理 · 甲方培训负责人',    '需求清单、权限矩阵、验收口径'],
        ['原型设计',   '第 2 周',   '产品经理 · UI 设计师',         '主要页面原型，双方评审确认'],
        ['功能开发',   '第 3–6 周', '前端 · 后端 · 数据工程师',     '七大功能模块可联调版本'],
        ['测试与修复', '第 7 周',   '测试工程师 · 甲方内测团队',    '测试记录、修复清单、权限与报表验证'],
        ['上线验收',   '第 8 周',   '产品经理 · 实施顾问',          '验收版本 · 管理员培训 · 操作手册'],
    ],
    col_widths_mm=[22, 16, 56, 61],
)

# ── 五、报价方案 ──────────────────────────────────────────────────────────────
doc.add_heading('五、报价方案', level=1)

doc.add_heading('5.1  席位规划说明', level=2)
doc.add_body(
    '贵司全国共 2,500 名员工，其中 500 名办公室人员有固定电脑，2,000 名为一线员工'
    '且流动率较高。针对两类人群的使用场景差异，我方建议采用"分层席位"模型——'
    '分别为两类人群提供匹配其实际需求的访问权限与账号管理方式，避免为一线员工'
    '支付不必要的高功能溢价，同时保障管理侧的完整数据视图。'
)

_fancy_table(
    raw_doc,
    headers=['席位类型', '适用人群', '访问方式', '核心功能', '账号管理'],
    data_rows=[
        ['A 类 · 全功能席位',
         '办公室人员\n500 人',
         'PC + 移动端',
         '全部功能，含 AI 辅助、高级报表、完整管理权限',
         '固定命名账号，长期有效'],
        ['B 类 · 流动席位池',
         '一线员工\n2,000 人额度',
         '移动端为主',
         '视频培训、课程完成、考试、积分、基础台账导出',
         '支持快速回收再分配，离职账号可复用给新员工，无需额外付费'],
    ],
    col_widths_mm=[28, 22, 22, 55, 28],
)

doc.add_note(
    '流动席位池的核心价值在于：当一线员工离职后，该账号可被管理员一键回收并重新'
    '分配给新入职员工，而不占用新的席位额度。这意味着无论年流动人数是 500 人还是'
    '1,000 人，席位费用始终固定——贵司只需为"2,000 个同时在册的一线学习席位"付费。',
    label='流动席位说明  ',
)

doc.add_heading('5.2  费用明细', level=2)
doc.add_body(
    '以下报价基于贵司 2,500 人规模（500 名全功能席位 + 2,000 名流动席位池）制定，'
    '含全功能部署、实施落地及上线后 1 年技术支持与免费维护服务。'
)

_fancy_table(
    raw_doc,
    headers=['费用项目', '说明', '规格', '单价', '金额（元）'],
    data_rows=[
        ['A 类席位年费',
         '办公室人员，全功能访问',
         '500 人/年',
         '¥160 / 人 / 年',
         '¥ 80,000'],
        ['B 类席位池年费',
         '一线员工，流动席位池',
         '2,000 额度 / 年',
         '¥30 / 额度 / 年',
         '¥ 60,000'],
        ['实施部署费（一次性）',
         '平台配置 · 组织架构导入 · 权限初始化 · 管理员培训（含 1 年免费维护）',
         '1 次',
         '—',
         '¥ 38,000'],
        ['年度技术支持',
         'Bug 修复 · 版本升级 · 工单响应（1 个工作日）',
         '含于席位年费',
         '—',
         '含  税'],
    ],
    col_widths_mm=[32, 52, 22, 26, 23],
)

doc.add_spacer(1)

# ── 合计高亮块（圆角卡片风格） ────────────────────────────────────────────────
total_tbl = raw_doc.add_table(2, 2)
total_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
_apply_card_style(total_tbl)

# Override insideH separator to use primary green for the one divider between rows
tblPr_t = _get_or_add_tblPr(total_tbl)
tblBdr_t = tblPr_t.find(qn('w:tblBorders'))
if tblBdr_t is not None:
    ih_t = tblBdr_t.find(qn('w:insideH'))
    if ih_t is not None:
        ih_t.set(qn('w:color'), BT.PRIMARY_500_HEX.lstrip('#'))
        ih_t.set(qn('w:sz'), '6')

r0 = total_tbl.rows[0]
for c in r0.cells:
    _cell_shading(c, BT.NEUTRAL_100_HEX)
r0.cells[0].width = Mm(100)
r0.cells[1].width = Mm(55)

lp0 = r0.cells[0].paragraphs[0]
lp0.paragraph_format.space_before = Pt(8)
lp0.paragraph_format.space_after  = Pt(4)
lp0.paragraph_format.left_indent  = Pt(10)
r_l0 = lp0.add_run('第一年合计（含实施费）')
_set_font(r_l0, size_pt=11, bold=True, color_hex=BT.NEUTRAL_900_HEX)
_set_run_font(r_l0)

rp0 = r0.cells[1].paragraphs[0]
rp0.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
rp0.paragraph_format.space_before = Pt(8)
rp0.paragraph_format.space_after  = Pt(4)
rp0.paragraph_format.right_indent = Pt(10)
r_r0 = rp0.add_run('¥ 178,000')
_set_font(r_r0, size_pt=14, bold=True, color_hex=BT.PRIMARY_500_HEX)
_set_run_font(r_r0)

r1 = total_tbl.rows[1]
r1.cells[0].width = Mm(100)
r1.cells[1].width = Mm(55)

lp1 = r1.cells[0].paragraphs[0]
lp1.paragraph_format.space_before = Pt(4)
lp1.paragraph_format.space_after  = Pt(8)
lp1.paragraph_format.left_indent  = Pt(10)
r_l1 = lp1.add_run('第二年起年度续费')
_set_font(r_l1, size_pt=10.5, color_hex=BT.NEUTRAL_700_HEX)
_set_run_font(r_l1)

rp1 = r1.cells[1].paragraphs[0]
rp1.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
rp1.paragraph_format.space_before = Pt(4)
rp1.paragraph_format.space_after  = Pt(8)
rp1.paragraph_format.right_indent = Pt(10)
r_r1 = rp1.add_run('¥ 140,000 / 年')
_set_font(r_r1, size_pt=11, bold=True, color_hex=BT.NEUTRAL_700_HEX)
_set_run_font(r_r1)

doc.add_heading('5.3  报价说明', level=2)
_fancy_table(
    raw_doc,
    headers=['条款', '说明'],
    data_rows=[
        ['价格说明',     '以上报价均为不含税价，开具增值税专用发票（税率 6%）'],
        ['有效期',       '本报价自文件出具之日起 60 个自然日内有效'],
        ['长期合作优惠', '若合同期签署 2 年，第 2 年席位年费享九五折优惠（¥133,000 / 年）'],
        ['席位调整',     '合同期内如办公室人员规模发生变化，可按实际人数等比调整 A 类席位费用'],
        ['增项服务',     'API 对接贵司 HR 系统、特种作业证书管理模块等可作为增项单独报价'],
    ],
    col_widths_mm=[35, 120],
)

doc.add_spacer(1)
doc.add_note(
    '我们已为 Vesuvius 中国区预留了专项资源，欢迎随时安排平台演示或技术对话。'
    '期待与贵团队开始这段合作。',
    label='联系我们  ',
)
doc.add_spacer(2)
doc.add_signature_block(company=BT.BRAND_FULL_CN, show_seal=True)

# ── 后处理 ────────────────────────────────────────────────────────────────────
_set_doc_font(raw_doc)

OUTPUT  = os.path.join(_DOCS, 'Vesuvius_培训平台增强建议书_v3_final.docx')
V1_DOCX = os.path.join(_DOCS, 'Vesuvius_培训平台增强建议书_v1.0.docx')

doc.save(OUTPUT)
_inject_saved_header(OUTPUT, V1_DOCX)

print(f'✓ Saved: {OUTPUT}')
