"""
BrandDocx — Brand-consistent DOCX builder for 未来方舟 / ArktechX.

Usage:
    from scripts.docx_builder import BrandDocx

    doc = BrandDocx()
    doc.add_title("在职证明", subtitle="EMPLOYMENT VERIFICATION LETTER")
    doc.add_body("This is to certify that ...")
    doc.add_note("本证明自盖章之日起 90 日内有效。")
    doc.save("output.docx")
"""

import io, zipfile, os
from typing import Optional, List, Tuple
import scripts.brand_tokens as BT

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def _rgb(hex_color: str) -> RGBColor:
    r, g, b = int(hex_color[1:3], 16), int(hex_color[3:5], 16), int(hex_color[5:7], 16)
    return RGBColor(r, g, b)

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_font(run, size_pt=None, bold=False, color_hex=None, italic=False):
    run.font.name = BT.FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BT.FONT_CN)
    if size_pt:  run.font.size    = Pt(size_pt)
    if bold:     run.font.bold    = True
    if italic:   run.font.italic  = True
    if color_hex: run.font.color.rgb = _rgb(color_hex)

def _cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    old  = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def _cell_borders(cell, sides: dict = None, **kwargs):
    merged = dict(sides or {})
    merged.update(kwargs)
    tcPr = cell._tc.get_or_add_tcPr()
    old  = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tc_borders = OxmlElement("w:tcBorders")
    for side, v in merged.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   v.get("val",   "single"))
        el.set(qn("w:sz"),    str(v.get("sz", 4)))
        el.set(qn("w:color"), v.get("color", BT.NEUTRAL_200_HEX.lstrip("#")))
        tc_borders.append(el)
    tcPr.append(tc_borders)

def _set_table_width(tbl, width_mm: float):
    """Fix total table width and strip the default TableGrid style so only
    explicit _cell_borders calls determine border rendering."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    # Remove python-docx's default w:tblStyle (usually TableGrid) so style-level
    # borders don't bleed through our cell-level border settings.
    old_style = tblPr.find(qn('w:tblStyle'))
    if old_style is not None:
        tblPr.remove(old_style)
    # Remove any direct tblBorders too
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'),    str(int(width_mm * 56.69)))  # mm → twips
    tblW.set(qn('w:type'), 'dxa')
    # Lock column widths so Word doesn't auto-resize and break cover layout
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

def _clear_doc_part(part):
    """Clear all tables and paragraphs from a header/footer part (makes it blank)."""
    for t in list(part.tables):
        t._tbl.getparent().remove(t._tbl)
    for p in list(part.paragraphs):
        p._p.getparent().remove(p._p)

def _set_row_height(row, height_mm: float, exact: bool = False):
    """Set minimum (or exact) row height. Uses same mm→twips factor as _set_table_width."""
    trPr     = row._tr.get_or_add_trPr()
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'),   str(int(height_mm * 56.69)))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')

def _set_cell_margins(cell, top=120, left=140, bottom=120, right=140):
    """Set per-cell padding. 120 twips ≈ 2.1 mm, 140 twips ≈ 2.5 mm."""
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")

def _repeat_table_header(row):
    """Mark a row as a repeating header (shows on every page when table spans pages)."""
    trPr      = row._tr.get_or_add_trPr()
    tblHeader = trPr.find(qn("w:tblHeader"))
    if tblHeader is None:
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    tblHeader.set(qn("w:val"), "true")

_NONE  = dict(val="none",   sz=0, color="FFFFFF")
_GREEN = dict(val="single", sz=6, color=BT.PRIMARY_500_HEX.lstrip("#"))
_GRAY  = dict(val="single", sz=4, color=BT.NEUTRAL_200_HEX.lstrip("#"))

def _para_left_bar(para, hex_color: str = BT.PRIMARY_500_HEX):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")    # 3pt
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)

def _para_shading(para, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    para._p.get_or_add_pPr().append(shd)

def _para_bottom_rule(para, hex_color: str = BT.PRIMARY_500_HEX):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)

# ── OOXML SVG injection (same as gen_employment_cert_en.py) ──────────────────

def _inject_svg(docx_path: str, svg_path: str, out_path: str):
    RELS_NS  = "http://schemas.openxmlformats.org/package/2006/relationships"
    IMG_TYPE = f"{RELS_NS[:-12]}/officeDocument/2006/relationships/image"
    A_NS     = "http://schemas.openxmlformats.org/drawingml/2006/main"
    R_NS     = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ASVG_NS  = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
    SVG_EXT  = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"

    with open(svg_path, "rb") as f:
        svg_bytes = f.read()

    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        hdrs  = sorted(n for n in names if n.startswith("word/header") and n.endswith(".xml"))
        if not hdrs:
            import shutil; shutil.copy(docx_path, out_path); return
        hdr_part  = hdrs[0]
        hdr_name  = os.path.basename(hdr_part)
        rels_part = f"word/_rels/{hdr_name}.rels"

        hdr_xml  = zin.read(hdr_part)
        rels_xml = zin.read(rels_part)

        rels_tree = etree.fromstring(rels_xml)
        all_rels  = rels_tree.findall(f"{{{RELS_NS}}}Relationship")
        used_nums = []
        for el in all_rels:
            try: used_nums.append(int(el.get("Id","rId0").replace("rId","")))
            except: pass
        svg_rId = f"rId{max(used_nums, default=0) + 1}"

        etree.SubElement(rels_tree, f"{{{RELS_NS}}}Relationship", attrib={
            "Id": svg_rId, "Type": IMG_TYPE, "Target": "media/logo_primary.svg"})
        new_rels = etree.tostring(rels_tree, xml_declaration=True,
                                  encoding="UTF-8", standalone=True)

        hdr_tree = etree.fromstring(hdr_xml)
        for blip in hdr_tree.findall(f".//{{{A_NS}}}blip"):
            extLst = blip.find(f"{{{A_NS}}}extLst")
            if extLst is None:
                extLst = etree.SubElement(blip, f"{{{A_NS}}}extLst")
            ext = etree.SubElement(extLst, f"{{{A_NS}}}ext", attrib={"uri": SVG_EXT})
            etree.SubElement(ext, f"{{{ASVG_NS}}}svgBlip",
                             attrib={f"{{{R_NS}}}embed": svg_rId})
            break
        new_hdr = etree.tostring(hdr_tree, xml_declaration=True,
                                 encoding="UTF-8", standalone=True)

        # Patch [Content_Types].xml to register SVG mime type
        CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
        ct_xml = zin.read("[Content_Types].xml")
        ct_tree = etree.fromstring(ct_xml)
        # Add <Default Extension="svg" ContentType="image/svg+xml"/> if missing
        existing_exts = {el.get("Extension","").lower()
                         for el in ct_tree.findall(f"{{{CT_NS}}}Default")}
        if "svg" not in existing_exts:
            etree.SubElement(ct_tree, f"{{{CT_NS}}}Default",
                             attrib={"Extension": "svg",
                                     "ContentType": "image/svg+xml"})
        new_ct = etree.tostring(ct_tree, xml_declaration=True,
                                encoding="UTF-8", standalone=True)

        with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == hdr_part:
                    zout.writestr(info, new_hdr)
                elif info.filename == rels_part:
                    zout.writestr(info, new_rels)
                elif info.filename == "[Content_Types].xml":
                    zout.writestr(info, new_ct)
                else:
                    zout.writestr(info, data)
            zout.writestr("word/media/logo_primary.svg", svg_bytes)


# ── BrandDocx class ───────────────────────────────────────────────────────────

class BrandDocx:
    """
    Brand-consistent DOCX builder for ArktechX documents.

    Example:
        doc = BrandDocx(doc_type="在职证明")
        doc.add_title("在职证明", subtitle="EMPLOYMENT VERIFICATION LETTER")
        doc.add_body("正文内容 ...")
        doc.save("out.docx")
    """

    def __init__(self, doc_type: str = ""):
        self._doc      = Document()
        self._doc_type = doc_type
        self._setup_page()
        self._setup_header()
        self._setup_footer()

    # ── Page ─────────────────────────────────────────────────────────────────

    def _setup_page(self):
        sec = self._doc.sections[0]
        sec.page_height     = Mm(297)
        sec.page_width      = Mm(210)
        sec.top_margin      = Mm(BT.PAGE_MARGIN_TOP_MM)
        sec.bottom_margin   = Mm(BT.PAGE_MARGIN_BOTTOM_MM)
        sec.left_margin     = Mm(BT.PAGE_MARGIN_LEFT_MM)
        sec.right_margin    = Mm(BT.PAGE_MARGIN_RIGHT_MM)
        sec.header_distance = Mm(10)
        sec.footer_distance = Mm(8)

    # ── Header ───────────────────────────────────────────────────────────────

    def _setup_header(self):
        sec = self._doc.sections[0]
        hdr = sec.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p._p.getparent().remove(p._p)

        # Use the actual brand PNG directly for consistent rendering
        with open(BT.LOGO_HORIZONTAL_PRIMARY_PNG, 'rb') as _f:
            logo_buf = io.BytesIO(_f.read())

        htbl = hdr.add_table(1, 2, width=Mm(155))
        htbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        lc = htbl.rows[0].cells[0]
        lc.width = Mm(110)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_borders(lc, top=_NONE, bottom=_GREEN, left=_NONE, right=_NONE)
        lp = lc.paragraphs[0]
        lp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        logo_run = lp.add_run()
        logo_run.add_picture(logo_buf, width=Mm(BT.LOGO_HORIZONTAL_ASPECT * 12))

        rc = htbl.rows[0].cells[1]
        rc.width = Mm(45)
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_borders(rc, top=_NONE, bottom=_GREEN, left=_NONE, right=_NONE)
        rp = rc.paragraphs[0]
        rp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(5)
        if self._doc_type:
            r = rp.add_run(self._doc_type.upper())
            _set_font(r, size_pt=8, color_hex=BT.NEUTRAL_400_HEX)

    def _make_logo_png(self) -> io.BytesIO:
        from PIL import Image, ImageDraw
        w_mm = BT.LOGO_HORIZONTAL_ASPECT * 12   # ≈ 47.5 mm
        w, h = int(w_mm * 11), int(12 * 11)     # ≈ 522 × 132 px
        img  = Image.new("RGB", (w, h), (255, 255, 255))
        d    = ImageDraw.Draw(img)
        mw   = int(w * 0.21)
        d.rectangle([0, 0, mw, h],      fill=BT.PRIMARY_500_RGB)
        d.rectangle([mw+1, h-3, w, h],  fill=BT.PRIMARY_500_RGB)
        buf  = io.BytesIO()
        img.save(buf, "PNG", dpi=(300, 300))
        buf.seek(0)
        return buf

    # ── Footer ───────────────────────────────────────────────────────────────

    def _setup_footer(self):
        sec = self._doc.sections[0]
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        for p in ftr.paragraphs:
            p.clear()

        ftbl = ftr.add_table(1, 2, width=Mm(155))
        ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in ftbl.rows:
            for c in row.cells:
                _cell_borders(c, top=_GRAY, bottom=_NONE, left=_NONE, right=_NONE)

        lc = ftbl.rows[0].cells[0]
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3)
        r = lp.add_run(BT.BRAND_FULL_CN)
        _set_font(r, size_pt=8, color_hex=BT.NEUTRAL_400_HEX)

        rc = ftbl.rows[0].cells[1]
        rp = rc.paragraphs[0]
        rp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(3)
        r2 = rp.add_run("本文件由人力资源部门出具，仅供证明使用")
        _set_font(r2, size_pt=7.5, color_hex=BT.NEUTRAL_400_HEX, italic=True)

        for p in list(ftr.paragraphs):
            if not p.text.strip():
                try: p._p.getparent().remove(p._p)
                except: pass

    # ── Public content methods ────────────────────────────────────────────────

    def add_cover_page(
        self,
        title: str,
        subtitle: str = "",
        doc_type: str = "",
        version: str = "",
        date: str = "",
        client: str = "",
        client_logo_path: str = "",
        illustration_path: str = "default",
        classification: str = "CONFIDENTIAL",
    ):
        """
        Full-page cover with hidden first-page header/footer.

        Layout (3-row skeleton table, no borders):
          Row 0 — our brand logo (left) + doc_type label (right) + green rule
          Row 1 — 2-col inner table: title zone (left ~100mm) | illustration (right ~55mm)
          Row 2 — right-aligned meta block (client / version / date / classification)
                  + green rule + our brand name

        Parameters
        ----------
        title            : main document title
        subtitle         : English subtitle (optional)
        doc_type         : label shown top-left, e.g. "Solution Proposal"
        version          : e.g. "V4.0"
        date             : e.g. "2026-06-11"
        client           : client company name
        client_logo_path : path to client logo PNG (optional; skip if empty/missing)
        illustration_path: "default" → assets/装饰性元素/3.png
                           ""        → no illustration
                           <path>    → custom image
        classification   : shown bottom-right, e.g. "CONFIDENTIAL" or "机密文件"
        """
        # ── 1. Hide first-page header/footer ─────────────────────────────────
        sec = self._doc.sections[0]
        sec.different_first_page_header_footer = True
        _clear_doc_part(sec.first_page_header)
        _clear_doc_part(sec.first_page_footer)

        # ── 2. Resolve illustration path ──────────────────────────────────────
        if illustration_path == "default":
            _illust = BT.DECO_3_PNG if os.path.exists(BT.DECO_3_PNG) else ""
        else:
            _illust = illustration_path

        # ── 3. Skeleton table (3 rows × 1 col) ───────────────────────────────
        PAGE_W = 155   # usable width mm (matches header/footer tables)
        cover  = self._doc.add_table(rows=3, cols=1)
        cover.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply width + fixed layout WITHOUT stripping tblStyle via _set_table_width,
        # because cover tables must not inherit data-table border rules.
        tblPr = cover._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            cover._tbl.insert(0, tblPr)
        for tag in ('w:tblStyle', 'w:tblBorders'):
            old = tblPr.find(qn(tag))
            if old is not None:
                tblPr.remove(old)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'),    str(int(PAGE_W * 56.69)))
        tblW.set(qn('w:type'), 'dxa')
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')

        for row in cover.rows:
            for c in row.cells:
                c.width = Mm(PAGE_W)
                _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

        _set_row_height(cover.rows[0], 28)
        _set_row_height(cover.rows[2], 65)

        # ── 4. Row 0 — brand logo + doc_type + green rule ─────────────────────
        top_cell = cover.rows[0].cells[0]
        top_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        logo_p = top_cell.paragraphs[0]
        logo_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after  = Pt(0)
        if os.path.exists(BT.LOGO_HORIZONTAL_PRIMARY_PNG):
            with open(BT.LOGO_HORIZONTAL_PRIMARY_PNG, 'rb') as _f:
                logo_p.add_run().add_picture(
                    io.BytesIO(_f.read()),
                    width=Mm(BT.LOGO_HORIZONTAL_ASPECT * 9),   # ~35 mm wide
                )

        if doc_type:
            type_p = top_cell.add_paragraph()
            type_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            type_p.paragraph_format.space_before = Pt(6)
            type_p.paragraph_format.space_after  = Pt(0)
            r = type_p.add_run(doc_type.upper())
            _set_font(r, size_pt=8, bold=True, color_hex=BT.PRIMARY_500_HEX)

        rule_p = top_cell.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(6)
        rule_p.paragraph_format.space_after  = Pt(0)
        _para_bottom_rule(rule_p, BT.PRIMARY_500_HEX)

        # ── 5. Row 1 — inner 2-col table (title left | illustration right) ────
        mid_cell = cover.rows[1].cells[0]
        mid_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_margins(mid_cell, top=200, left=0, bottom=200, right=0)

        LEFT_W  = 100
        RIGHT_W = PAGE_W - LEFT_W   # 55 mm

        inner = mid_cell.add_table(rows=1, cols=2)
        inner.alignment = WD_TABLE_ALIGNMENT.LEFT
        for c in inner.rows[0].cells:
            _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

        # Remove style/borders from inner table
        itblPr = inner._tbl.find(qn('w:tblPr'))
        if itblPr is None:
            itblPr = OxmlElement('w:tblPr')
            inner._tbl.insert(0, itblPr)
        for tag in ('w:tblStyle', 'w:tblBorders'):
            old = itblPr.find(qn(tag))
            if old is not None:
                itblPr.remove(old)

        left_c  = inner.rows[0].cells[0]
        right_c = inner.rows[0].cells[1]
        left_c.width  = Mm(LEFT_W)
        right_c.width = Mm(RIGHT_W)
        left_c.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
        right_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Client logo (optional)
        if client_logo_path and os.path.exists(client_logo_path):
            cl_p = left_c.paragraphs[0]
            cl_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            cl_p.paragraph_format.space_before = Pt(0)
            cl_p.paragraph_format.space_after  = Pt(14)
            with open(client_logo_path, 'rb') as _f:
                cl_p.add_run().add_picture(io.BytesIO(_f.read()), width=Mm(44))
            title_para = left_c.add_paragraph()
        else:
            title_para = left_c.paragraphs[0]

        # Main title
        title_para.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after  = Pt(8)
        r = title_para.add_run(title)
        _set_font(r, size_pt=26, bold=True, color_hex=BT.NEUTRAL_900_HEX)

        # Subtitle
        if subtitle:
            sub_p = left_c.add_paragraph()
            sub_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            sub_p.paragraph_format.space_before = Pt(0)
            sub_p.paragraph_format.space_after  = Pt(14)
            r = sub_p.add_run(subtitle)
            _set_font(r, size_pt=10.5, color_hex=BT.NEUTRAL_400_HEX)

        # Green rule under title block
        title_rule = left_c.add_paragraph()
        title_rule.paragraph_format.space_before = Pt(4)
        title_rule.paragraph_format.space_after  = Pt(0)
        _para_bottom_rule(title_rule, BT.PRIMARY_500_HEX)

        # Illustration (right column)
        if _illust and os.path.exists(_illust):
            ill_p = right_c.paragraphs[0]
            ill_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
            ill_p.paragraph_format.space_before = Pt(0)
            ill_p.paragraph_format.space_after  = Pt(0)
            with open(_illust, 'rb') as _f:
                ill_p.add_run().add_picture(
                    io.BytesIO(_f.read()), width=Mm(RIGHT_W - 3)
                )

        # ── 6. Row 2 — meta info block ────────────────────────────────────────
        bot_cell = cover.rows[2].cells[0]
        bot_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        _set_cell_margins(bot_cell, top=0, left=0, bottom=120, right=0)

        meta_items = [
            ("客户",  client),
            ("版本",  version),
            ("日期",  date),
            ("密级",  classification),
        ]
        for label, value in meta_items:
            if not value:
                continue
            mp = bot_cell.add_paragraph()
            mp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
            mp.paragraph_format.space_before = Pt(2)
            mp.paragraph_format.space_after  = Pt(2)
            r1 = mp.add_run(f"{label}  ")
            _set_font(r1, size_pt=8.5, bold=True, color_hex=BT.NEUTRAL_400_HEX)
            r2 = mp.add_run(str(value))
            _set_font(r2, size_pt=8.5, color_hex=BT.NEUTRAL_700_HEX)

        # Bottom rule + our brand name
        br_p = bot_cell.add_paragraph()
        br_p.paragraph_format.space_before = Pt(10)
        br_p.paragraph_format.space_after  = Pt(4)
        _para_bottom_rule(br_p, BT.NEUTRAL_200_HEX)

        brand_p = bot_cell.add_paragraph()
        brand_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.RIGHT
        brand_p.paragraph_format.space_before = Pt(0)
        brand_p.paragraph_format.space_after  = Pt(0)
        r = brand_p.add_run(BT.BRAND_FULL_CN)
        _set_font(r, size_pt=8.5, color_hex=BT.NEUTRAL_400_HEX)

        # ── 7. Page break to start content on page 2 ─────────────────────────
        pb = self._doc.add_paragraph()
        pb.paragraph_format.space_before = Pt(0)
        pb.paragraph_format.space_after  = Pt(0)
        pb.add_run().add_break(WD_BREAK.PAGE)

    def add_spacer(self, lines: int = 1):
        for _ in range(lines):
            self._doc.add_paragraph()

    def add_title(self, main: str, subtitle: str = ""):
        """Large centered title with green bottom rule."""
        tp = self._doc.add_paragraph()
        tp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(2)
        r = tp.add_run(main)
        _set_font(r, size_pt=20, bold=True, color_hex=BT.NEUTRAL_900_HEX)

        if subtitle:
            sp = self._doc.add_paragraph()
            sp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(12)
            rs = sp.add_run(subtitle)
            _set_font(rs, size_pt=9, color_hex=BT.NEUTRAL_400_HEX)

        rule_p = self._doc.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(0)
        rule_p.paragraph_format.space_after  = Pt(16)
        _para_bottom_rule(rule_p)

    def add_heading(self, text: str, level: int = 1):
        """H1–H3 with brand colors."""
        p = self._doc.add_paragraph()
        if level == 1:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text)
            _set_font(r, size_pt=BT.FONT_H1_PT, bold=True,
                      color_hex=BT.FONT_H1_COLOR)
        elif level == 2:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            _set_font(r, size_pt=BT.FONT_H2_PT, bold=True,
                      color_hex=BT.FONT_H2_COLOR)
        else:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            _set_font(r, size_pt=BT.FONT_H3_PT, bold=True,
                      color_hex=BT.FONT_H3_COLOR)
        return p

    def add_body(self, text: str, indent: bool = True, line_spacing: float = 1.5):
        """Standard body paragraph."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        if indent:
            p.paragraph_format.first_line_indent = Pt(22)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing      = line_spacing
        r = p.add_run(text)
        _set_font(r, size_pt=BT.FONT_BODY_PT, color_hex=BT.NEUTRAL_700_HEX)
        return p

    def add_note(self, text: str, label: str = "注："):
        """Left accent bar callout block."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(8)
        p.paragraph_format.right_indent = Pt(8)
        _para_left_bar(p)
        _para_shading(p, BT.NEUTRAL_100_HEX)
        if label:
            r1 = p.add_run(label)
            _set_font(r1, size_pt=9.5, bold=True,
                      color_hex=BT.PRIMARY_500_HEX)
        r2 = p.add_run(text)
        _set_font(r2, size_pt=9.5, color_hex=BT.NEUTRAL_400_HEX)
        return p

    def add_info_table(self,
                       rows: List[Tuple[str, str]],
                       label_width_mm: float = 38,
                       value_width_mm: float = 117,
                       table_width_mm: float = 155):
        """
        Key-value info table (e.g. for employment cert fields).
        rows: [("姓名", "_____"), ("部门", "_____"), ...]
        """
        tbl = self._doc.add_table(rows=len(rows), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_width(tbl, table_width_mm)

        _ROW_BOTTOM = dict(val="single", sz=4, color=BT.NEUTRAL_200_HEX.lstrip("#"))

        for i, (label, value) in enumerate(rows):
            # ── Label cell ────────────────────────────────────────────────────
            lc = tbl.rows[i].cells[0]
            lc.width = Mm(label_width_mm)
            lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_shading(lc, BT.NEUTRAL_100_HEX)
            _set_cell_margins(lc, top=130, left=130, bottom=130, right=120)
            _cell_borders(lc, top=_NONE, bottom=_ROW_BOTTOM, left=_NONE, right=_NONE)
            lp = lc.paragraphs[0]
            lp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after  = Pt(0)
            r = lp.add_run(label)
            _set_font(r, size_pt=9.5, bold=True, color_hex=BT.NEUTRAL_700_HEX)

            # ── Value cell ────────────────────────────────────────────────────
            vc = tbl.rows[i].cells[1]
            vc.width = Mm(value_width_mm)
            vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(vc, top=130, left=160, bottom=130, right=130)
            _cell_borders(vc, top=_NONE, bottom=_ROW_BOTTOM, left=_NONE, right=_NONE)
            vp = vc.paragraphs[0]
            vp.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
            vp.paragraph_format.space_before = Pt(0)
            vp.paragraph_format.space_after  = Pt(0)
            vr = vp.add_run(str(value))
            _set_font(vr, size_pt=9.8, color_hex=BT.NEUTRAL_700_HEX)

        return tbl

    def add_data_table(self,
                       headers: List[str],
                       data_rows: List[List[str]],
                       col_widths_mm: Optional[List[float]] = None,
                       table_width_mm: float = 155,
                       first_col_bold: bool = False):
        """
        Proposal-style data table: horizontal rules only, no outer box, brand-color header.
        headers: ["列名1", "列名2", ...]
        data_rows: [["a","b"], ["c","d"], ...]
        """
        _HDR_BOTTOM = dict(val="single", sz=6, color=BT.PRIMARY_500_HEX.lstrip("#"))
        _ROW_BOTTOM = dict(val="single", sz=4, color=BT.NEUTRAL_200_HEX.lstrip("#"))

        n_cols = len(headers)
        tbl = self._doc.add_table(rows=1 + len(data_rows), cols=n_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_width(tbl, table_width_mm)

        # Header row
        hrow = tbl.rows[0]
        _repeat_table_header(hrow)
        for j, hdr_text in enumerate(headers):
            c = hrow.cells[j]
            if col_widths_mm:
                c.width = Mm(col_widths_mm[j])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_shading(c, BT.TABLE_HEADER_BG)
            _set_cell_margins(c, top=150, left=130, bottom=150, right=130)
            _cell_borders(c, top=_NONE, bottom=_HDR_BOTTOM, left=_NONE, right=_NONE)
            p = c.paragraphs[0]
            p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            r = p.add_run(hdr_text)
            _set_font(r, size_pt=9.5, bold=True, color_hex=BT.TABLE_HEADER_FG)

        # Data rows
        for i, row_data in enumerate(data_rows):
            tr = tbl.rows[i + 1]
            for j, cell_text in enumerate(row_data):
                c = tr.cells[j]
                if col_widths_mm:
                    c.width = Mm(col_widths_mm[j])
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i % 2 == 1:
                    _cell_shading(c, BT.TABLE_STRIPE_BG)
                _set_cell_margins(c, top=130, left=130, bottom=130, right=130)
                _cell_borders(c, top=_NONE, bottom=_ROW_BOTTOM, left=_NONE, right=_NONE)
                p = c.paragraphs[0]
                p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.line_spacing = 1.2
                r = p.add_run(str(cell_text))
                _set_font(r, size_pt=9.5,
                          bold=(first_col_bold and j == 0),
                          color_hex=BT.NEUTRAL_700_HEX)

        return tbl

    def add_signature_block(self,
                            company: str = BT.BRAND_FULL_CN,
                            show_seal: bool = True):
        """Standard 2-column signature/seal block."""
        self.add_spacer(2)

        sig = self._doc.add_table(rows=3, cols=2)
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in sig.rows:
            for c in row.cells:
                _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

        def _sig_row(i, ll, lv, rl="", rv=""):
            for c in sig.rows[i].cells:
                c.width = Mm(77)
            lp = sig.rows[i].cells[0].paragraphs[0]
            rp = sig.rows[i].cells[1].paragraphs[0]
            for p in (lp, rp):
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after  = Pt(5)
            for text, p in [(ll, lp), (lv, lp), (rl, rp), (rv, rp)]:
                r = p.add_run(text)
                _set_font(r, size_pt=10.5, color_hex=BT.NEUTRAL_700_HEX)

        _sig_row(0, "公司名称：", company)
        _sig_row(1, "授权签字：", "＿" * 12,
                    "公司盖章：", "（公章）" if show_seal else "")
        _sig_row(2, "出具日期：", "     年      月      日")

    def add_signature_block_en(self, company: str = BT.BRAND_FULL_EN):
        """English version signature block."""
        self.add_spacer(2)
        sig = self._doc.add_table(rows=3, cols=2)
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in sig.rows:
            for c in row.cells:
                _cell_borders(c, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

        def _row(i, ll, lv, rl="", rv=""):
            for c in sig.rows[i].cells:
                c.width = Mm(77)
            for p_idx, (label, val) in enumerate([(ll, lv), (rl, rv)]):
                p = sig.rows[i].cells[p_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after  = Pt(5)
                for t in (label, val):
                    r = p.add_run(t)
                    _set_font(r, size_pt=10.5, color_hex=BT.NEUTRAL_700_HEX)

        _row(0, "Company: ",  company)
        _row(1, "Authorized Signatory: ", "_" * 18,
                "Company Seal: ", "(Official Seal)")
        _row(2, "Date of Issuance: ",
                "________  Year    ______  Month    ______  Day")

    # ── Save ─────────────────────────────────────────────────────────────────

    def save(self, output_path: str, inject_svg: bool = False):
        """
        Save to output_path.
        inject_svg is False by default — the header already uses the brand PNG.
        Pass inject_svg=True only if SVG vector overlay is explicitly needed.
        """
        if inject_svg and os.path.exists(BT.LOGO_HORIZONTAL_PRIMARY_SVG):
            tmp = output_path + ".tmp.docx"
            self._doc.save(tmp)
            try:
                _inject_svg(tmp, BT.LOGO_HORIZONTAL_PRIMARY_SVG, output_path)
                os.remove(tmp)
            except Exception as e:
                os.rename(tmp, output_path)
                print(f"  Warning: SVG injection failed ({e}), saved PNG fallback.")
        else:
            self._doc.save(output_path)

    @property
    def document(self):
        """Direct access to the underlying python-docx Document."""
        return self._doc
