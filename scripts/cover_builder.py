"""
Cover Builder — 统一封面生成器，支持 DOCX 和 PDF 格式自适应布局

该模块提供统一的封面数据结构和格式无关的布局算法，
确保在不同输出格式（DOCX、PDF/HTML）中保持一致的视觉效果。

Usage:
    from scripts.cover_builder import CoverConfig, create_docx_cover, create_html_cover

    config = CoverConfig(
        title="企业培训平台增强建议书",
        subtitle="Enterprise Training Platform Enhancement Proposal",
        doc_type="SOLUTION PROPOSAL",
        version="V1.0",
        date="2026-06-15",
        client="某某客户",
        layout_type="business"  # 可选: "business", "technical", "presentation"
    )

    # DOCX 格式
    create_docx_cover(doc, config)

    # HTML 格式（用于 PDF 生成）
    html_content = create_html_cover(config)
"""

import os
import io
from typing import Optional, List, Dict, Any, Literal
from dataclasses import dataclass, field
import scripts.brand_tokens as BT


@dataclass
class CoverConfig:
    """统一封面配置数据结构"""

    # 核心内容
    title: str
    subtitle: str = ""
    doc_type: str = ""

    # 版本和日期
    version: str = ""
    date: str = ""

    # 客户信息
    client: str = ""
    client_logo_path: str = ""

    # 视觉元素
    illustration_path: str = "default"  # "default", "", 或自定义路径
    layout_type: Literal["business", "technical", "presentation"] = "business"

    # 机密级别
    classification: str = "CONFIDENTIAL"

    # 品牌元素
    brand_logo_path: str = field(default_factory=lambda: BT.LOGO_HORIZONTAL_PRIMARY_PNG)
    brand_name: str = field(default_factory=lambda: BT.BRAND_FULL_CN)

    # 布局控制
    show_decoration: bool = True
    show_client_logo: bool = True

    def __post_init__(self):
        """处理默认值和路径解析"""
        if self.illustration_path == "default":
            self.illustration_path = BT.DECO_3_PNG if os.path.exists(BT.DECO_3_PNG) else ""

        # 验证品牌 logo 路径
        if not os.path.exists(self.brand_logo_path):
            self.brand_logo_path = ""

        # 验证客户 logo 路径
        if self.client_logo_path and not os.path.exists(self.client_logo_path):
            self.client_logo_path = ""

    @property
    def has_illustration(self) -> bool:
        """是否有装饰插图"""
        return bool(self.illustration_path and os.path.exists(self.illustration_path))

    @property
    def has_brand_logo(self) -> bool:
        """是否有品牌 logo"""
        return bool(self.brand_logo_path and os.path.exists(self.brand_logo_path))

    @property
    def has_client_logo(self) -> bool:
        """是否有客户 logo"""
        return bool(self.client_logo_path and os.path.exists(self.client_logo_path) and self.show_client_logo)

    def get_meta_items(self) -> List[tuple]:
        """获取元信息项目列表（显示在封面底部）"""
        items = []
        if self.client:
            items.append(("客户", self.client))
        if self.version:
            items.append(("版本", self.version))
        if self.date:
            items.append(("日期", self.date))
        if self.classification:
            items.append(("密级", self.classification))
        return items


class LayoutCalculator:
    """布局计算器 - 提供格式无关的布局算法"""

    # 标准页面尺寸（mm）
    A4_WIDTH = 210
    A4_HEIGHT = 297

    # 标准边距（mm）
    MARGIN_LR = 31.7  # 左右边距
    MARGIN_TB = 25.4  # 上下边距

    @classmethod
    def get_content_area(cls) -> dict:
        """获取内容区域尺寸"""
        return {
            'width': cls.A4_WIDTH - 2 * cls.MARGIN_LR,    # 146.6 mm
            'height': cls.A4_HEIGHT - 2 * cls.MARGIN_TB   # 246.2 mm
        }

    @classmethod
    def calculate_layout(cls, config: CoverConfig) -> dict:
        """计算封面布局参数"""
        content = cls.get_content_area()

        # 根据布局类型调整比例
        layout_ratios = {
            'business': {'header': 0.18, 'content': 0.62, 'footer': 0.20},
            'technical': {'header': 0.15, 'content': 0.70, 'footer': 0.15},
            'presentation': {'header': 0.12, 'content': 0.76, 'footer': 0.12}
        }

        ratios = layout_ratios.get(config.layout_type, layout_ratios['business'])

        return {
            'content_width': content['width'],
            'content_height': content['height'],
            'header_height': content['height'] * ratios['header'],
            'main_height': content['height'] * ratios['content'],
            'footer_height': content['height'] * ratios['footer'],
            'logo_width': cls._calculate_logo_width(config),
            'title_font_size': cls._calculate_title_font_size(config),
            'has_two_column': config.has_illustration and len(config.title) > 20
        }

    @classmethod
    def _calculate_logo_width(cls, config: CoverConfig) -> float:
        """计算 logo 最佳宽度"""
        if config.layout_type == "presentation":
            return 45  # 演示文稿用较大 logo
        elif config.layout_type == "technical":
            return 35  # 技术文档用中等 logo
        else:
            return 40  # 商务文档用标准 logo

    @classmethod
    def _calculate_title_font_size(cls, config: CoverConfig) -> dict:
        """根据标题长度和布局类型计算字体大小"""
        title_len = len(config.title)

        # 基础字号
        base_sizes = {
            'business': {'title': 28, 'subtitle': 12},
            'technical': {'title': 26, 'subtitle': 11},
            'presentation': {'title': 32, 'subtitle': 14}
        }

        sizes = base_sizes.get(config.layout_type, base_sizes['business'])

        # 长标题缩小字号
        if title_len > 30:
            sizes['title'] -= 4
        elif title_len > 20:
            sizes['title'] -= 2

        return sizes


def create_docx_cover(doc, config: CoverConfig):
    """创建 DOCX 格式封面（优化版）"""
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from scripts.docx_builder import (
        _set_font, _cell_shading, _cell_borders, _clear_doc_part,
        _set_table_width, _set_row_height, _para_bottom_rule,
        _NONE, _GREEN, _GRAY
    )

    # 计算自适应布局
    layout = LayoutCalculator.calculate_layout(config)

    # 获取实际的 docx.Document 对象
    raw_doc = doc._doc if hasattr(doc, '_doc') else doc

    # 隐藏首页页眉页脚
    sec = raw_doc.sections[0]
    sec.different_first_page_header_footer = True
    _clear_doc_part(sec.first_page_header)
    _clear_doc_part(sec.first_page_footer)

    # 创建主表格（3行结构）
    cover = raw_doc.add_table(rows=3, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格宽度和布局
    _set_table_width(cover, layout['content_width'])

    # 清除所有边框
    for row in cover.rows:
        for cell in row.cells:
            cell.width = Mm(layout['content_width'])
            _cell_borders(cell, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)

    # 设置行高
    _set_row_height(cover.rows[0], layout['header_height'])
    _set_row_height(cover.rows[1], layout['main_height'])
    _set_row_height(cover.rows[2], layout['footer_height'])

    # === 头部区域 ===
    _create_docx_header_section(cover.rows[0].cells[0], config, layout)

    # === 主要内容区域 ===
    _create_docx_content_section(cover.rows[1].cells[0], config, layout)

    # === 底部区域 ===
    _create_docx_footer_section(cover.rows[2].cells[0], config, layout)

    # 添加分页符
    raw_doc.add_page_break()


def _create_docx_header_section(cell, config: CoverConfig, layout: dict):
    """创建 DOCX 头部区域"""
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from scripts.docx_builder import _set_font

    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 品牌 logo
    if config.has_brand_logo:
        logo_p = cell.paragraphs[0]
        logo_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after = Pt(8)

        with open(config.brand_logo_path, 'rb') as f:
            logo_p.add_run().add_picture(
                io.BytesIO(f.read()),
                width=Mm(layout['logo_width'])
            )

    # 文档类型标签
    if config.doc_type:
        type_p = cell.add_paragraph()
        type_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        type_p.paragraph_format.space_before = Pt(6)
        type_p.paragraph_format.space_after = Pt(8)

        type_run = type_p.add_run(config.doc_type.upper())
        _set_font(type_run, size_pt=10, bold=True, color_hex=BT.PRIMARY_500_HEX)

        # 添加绿色分隔线
        rule_p = cell.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(4)
        rule_p.paragraph_format.space_after = Pt(0)
        from scripts.docx_builder import _para_bottom_rule
        _para_bottom_rule(rule_p, BT.PRIMARY_500_HEX)


def _create_docx_content_section(cell, config: CoverConfig, layout: dict):
    """创建 DOCX 主要内容区域"""
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from scripts.docx_builder import _set_font, _set_table_width, _cell_borders, _NONE

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 如果有装饰图且标题较长，使用两列布局
    if layout['has_two_column']:
        content_table = cell.add_table(rows=1, cols=2)
        content_table.autofit = False
        _set_table_width(content_table, layout['content_width'])

        title_cell = content_table.rows[0].cells[0]
        image_cell = content_table.rows[0].cells[1]

        title_cell.width = Mm(layout['content_width'] * 0.68)
        image_cell.width = Mm(layout['content_width'] * 0.32)

        # 清除表格边框
        for cell_inner in content_table.rows[0].cells:
            _cell_borders(cell_inner, top=_NONE, bottom=_NONE, left=_NONE, right=_NONE)
    else:
        title_cell = cell
        image_cell = None

    # 主标题
    title_p = title_cell.paragraphs[0]
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    title_p.paragraph_format.line_spacing = 1.1
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)

    title_run = title_p.add_run(config.title)
    font_sizes = layout.get('title_font_size', LayoutCalculator._calculate_title_font_size(config))
    _set_font(title_run, size_pt=font_sizes['title'], bold=True, color_hex=BT.NEUTRAL_900_HEX)

    # 副标题
    if config.subtitle:
        subtitle_p = title_cell.add_paragraph()
        subtitle_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle_p.paragraph_format.space_before = Pt(4)
        subtitle_p.paragraph_format.space_after = Pt(12)

        subtitle_run = subtitle_p.add_run(config.subtitle)
        _set_font(subtitle_run, size_pt=font_sizes['subtitle'], color_hex=BT.NEUTRAL_400_HEX)
        # 英文副标题使用 Inter 字体
        subtitle_run.font.name = BT.FONT_EN

    # 装饰图片
    if config.has_illustration and image_cell:
        image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        img_p = image_cell.paragraphs[0]
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        with open(config.illustration_path, 'rb') as f:
            optimal_width = min(Mm(40), Mm(layout['content_width'] * 0.25))
            img_p.add_run().add_picture(io.BytesIO(f.read()), width=optimal_width)


def _create_docx_footer_section(cell, config: CoverConfig, layout: dict):
    """创建 DOCX 底部区域"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from scripts.docx_builder import _set_font, _para_bottom_rule

    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    # 元信息
    meta_items = config.get_meta_items()
    for i, (label, value) in enumerate(meta_items):
        if not value:
            continue

        meta_p = cell.add_paragraph()
        meta_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_p.paragraph_format.space_before = Pt(1)
        meta_p.paragraph_format.space_after = Pt(1)

        label_run = meta_p.add_run(f"{label}  ")
        _set_font(label_run, size_pt=9, bold=True, color_hex=BT.NEUTRAL_400_HEX)

        value_run = meta_p.add_run(str(value))
        _set_font(value_run, size_pt=9, color_hex=BT.NEUTRAL_700_HEX)

    # 底部分隔线和品牌名称
    if meta_items:
        rule_p = cell.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(8)
        rule_p.paragraph_format.space_after = Pt(4)
        _para_bottom_rule(rule_p, BT.PRIMARY_500_HEX)

        brand_p = cell.add_paragraph()
        brand_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        brand_p.paragraph_format.space_after = Pt(0)

        brand_run = brand_p.add_run(config.brand_name)
        _set_font(brand_run, size_pt=8, color_hex=BT.NEUTRAL_400_HEX)


def create_html_cover(config: CoverConfig) -> str:
    """创建 HTML 格式封面（用于 PDF 生成）"""
    layout = LayoutCalculator.calculate_layout(config)
    font_sizes = LayoutCalculator._calculate_title_font_size(config)

    # 构建 HTML 内容
    html = f"""
<!DOCTYPE html>
<html lang="zh-Hans">
<head>
    <meta charset="utf-8">
    <title>{config.title}</title>
    <link rel="stylesheet" href="../../../templates/html/brand.css">
    <style>
        .cover {{
            height: 297mm;
            display: flex;
            flex-direction: column;
            background: linear-gradient(180deg, #ffffff 0%, #ffffff 72%, #fcfffe 100%);
            {_get_decoration_background(config)}
        }}

        .cover__header {{
            height: {layout['header_height']:.1f}mm;
            padding: 12mm 31mm 0;
            display: flex;
            flex-direction: column;
            justify-content: flex-start;
        }}

        .cover__content {{
            height: {layout['main_height']:.1f}mm;
            padding: 0 31mm;
            display: flex;
            align-items: center;
            {_get_content_layout_styles(config, layout)}
        }}

        .cover__footer {{
            height: {layout['footer_height']:.1f}mm;
            padding: 0 31mm 12mm;
            display: flex;
            flex-direction: column;
            justify-content: flex-end;
            align-items: flex-end;
        }}

        .cover__logo {{
            width: {layout['logo_width']:.1f}mm;
            height: auto;
            margin-bottom: 6mm;
        }}

        .cover__doc-type {{
            font-size: 10pt;
            font-weight: 600;
            color: var(--color-primary-500);
            text-transform: uppercase;
            margin-bottom: 4mm;
            letter-spacing: 0.5px;
        }}

        .cover__title {{
            font-size: {font_sizes['title']}pt;
            font-weight: 700;
            color: var(--color-neutral-900);
            line-height: 1.1;
            margin-bottom: {8 if config.subtitle else 0}mm;
        }}

        .cover__subtitle {{
            font-size: {font_sizes['subtitle']}pt;
            color: var(--color-neutral-400);
            font-family: var(--font-en);
            font-weight: 400;
        }}

        .cover__illustration {{
            max-width: 45mm;
            max-height: 60mm;
            object-fit: contain;
        }}

        .cover__meta {{
            text-align: right;
            font-size: 9pt;
            line-height: 1.4;
            margin-bottom: 6mm;
        }}

        .cover__meta-item {{
            margin: 1mm 0;
        }}

        .cover__meta-label {{
            font-weight: 600;
            color: var(--color-neutral-400);
        }}

        .cover__meta-value {{
            color: var(--color-neutral-700);
        }}

        .cover__brand {{
            font-size: 8pt;
            color: var(--color-neutral-400);
            border-top: 1px solid var(--color-primary-500);
            padding-top: 4mm;
        }}

        .cover__rule {{
            height: 1px;
            background: var(--color-primary-500);
            margin: 4mm 0;
        }}
    </style>
</head>
<body>
    <div class="page cover">
        <div class="cover__header">
            {_render_header_section(config, layout)}
        </div>

        <div class="cover__content">
            {_render_content_section(config, layout)}
        </div>

        <div class="cover__footer">
            {_render_footer_section(config)}
        </div>
    </div>
</body>
</html>
"""

    return html.strip()


def _get_decoration_background(config: CoverConfig) -> str:
    """获取装饰背景样式"""
    if not config.has_illustration or config.layout_type == "technical":
        return ""

    return """
    background-image: radial-gradient(circle at 85% 82%, rgba(62, 201, 158, .06), transparent 18%);
    """


def _get_content_layout_styles(config: CoverConfig, layout: dict) -> str:
    """获取内容区域布局样式"""
    if layout['has_two_column']:
        return """
        justify-content: space-between;
        """
    else:
        return """
        justify-content: flex-start;
        """


def _render_header_section(config: CoverConfig, layout: dict) -> str:
    """渲染头部区域 HTML"""
    html_parts = []

    # 品牌 logo
    if config.has_brand_logo:
        # 转换为相对路径
        logo_rel_path = os.path.relpath(config.brand_logo_path,
                                       os.path.join(os.path.dirname(__file__), '..', 'projects'))
        logo_rel_path = logo_rel_path.replace('\\', '/')
        if not logo_rel_path.startswith('../'):
            logo_rel_path = '../' + logo_rel_path

        html_parts.append(f'<img src="{logo_rel_path}" class="cover__logo" alt="Brand Logo">')

    # 文档类型
    if config.doc_type:
        html_parts.append(f'<div class="cover__doc-type">{config.doc_type}</div>')
        html_parts.append('<div class="cover__rule"></div>')

    return '\n            '.join(html_parts)


def _render_content_section(config: CoverConfig, layout: dict) -> str:
    """渲染内容区域 HTML"""
    title_html = f'<h1 class="cover__title">{config.title}</h1>'
    if config.subtitle:
        title_html += f'<div class="cover__subtitle">{config.subtitle}</div>'

    if layout['has_two_column']:
        # 两列布局
        illustration_html = ""
        if config.has_illustration:
            img_rel_path = os.path.relpath(config.illustration_path,
                                         os.path.join(os.path.dirname(__file__), '..', 'projects'))
            img_rel_path = img_rel_path.replace('\\', '/')
            if not img_rel_path.startswith('../'):
                img_rel_path = '../' + img_rel_path

            illustration_html = f'<img src="{img_rel_path}" class="cover__illustration" alt="Illustration">'

        return f"""
            <div style="flex: 2;">
                {title_html}
            </div>
            <div style="flex: 1; text-align: right;">
                {illustration_html}
            </div>
        """
    else:
        # 单列布局
        return f'<div>{title_html}</div>'


def _render_footer_section(config: CoverConfig) -> str:
    """渲染底部区域 HTML"""
    html_parts = []

    # 元信息
    meta_items = config.get_meta_items()
    if meta_items:
        meta_html = []
        for label, value in meta_items:
            if value:
                meta_html.append(f'''
                    <div class="cover__meta-item">
                        <span class="cover__meta-label">{label}</span>
                        <span class="cover__meta-value">{value}</span>
                    </div>
                ''')

        html_parts.append(f'<div class="cover__meta">{"".join(meta_html)}</div>')
        html_parts.append(f'<div class="cover__brand">{config.brand_name}</div>')

    return '\n            '.join(html_parts)


# 向后兼容的便捷函数
def add_adaptive_cover_to_docx(doc, **kwargs):
    """向后兼容的 DOCX 封面添加函数"""
    config = CoverConfig(**kwargs)
    create_docx_cover(doc, config)


def generate_cover_html(**kwargs) -> str:
    """向后兼容的 HTML 封面生成函数"""
    config = CoverConfig(**kwargs)
    return create_html_cover(config)